# -*- coding: utf-8 -*-
"""
Created on Fri Dec 21 10:48:13 2018

@author: buddeja
"""

from tkinter import*
from PIL import Image,ImageTk
from tkinter import ttk
from specpy import *
import specpy
import numpy as np
from scipy.ndimage.filters import maximum_filter
from scipy.ndimage.morphology import generate_binary_structure, binary_erosion
from scipy.ndimage.filters import gaussian_filter
import matplotlib.pyplot as pp
import math
import os
from scipy.stats import stats
from skimage import img_as_float
from PIL import Image
from scipy import ndimage
from skimage import feature
import numpy.ma as ma
from skimage.feature import peak_local_max
import shutil
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg#, NavigationToolbar2TkAgg
from skimage import color
import matplotlib.cm as cm
from sklearn.neighbors import NearestNeighbors
from astropy.stats import RipleysKEstimator
from docx import Document
from docx.shared import Pt
from docx.shared import Length
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


import numpy as np
import csv
import pandas as pd
import io
from io import StringIO 
import scipy.stats
from scipy.spatial import distance
from astropy.stats import RipleysKEstimator
from sklearn.neighbors import NearestNeighbors
from functools import partial


root= Tk()
root.title("Imspector Control Interface")
root1 = Frame(root, width = 350, height = 350, bg = 'grey')
#root1.grid_propagate(0)
root1.grid()
#root.grid_rowconfigure(1, weight=1)
#root.grid_columnconfigure(1, weight=1)

foldername= 'testfolder'   #testfolder'#Type here the foldername
path = 'C:/Users/Abberior_admin/Desktop/GUI_20192611/GUI_pythoninterface/'    #'D:/current data/'#Type here the path


image_ID = Image.open('{}{}'.format(path,'Alphatubulin_TEST.tif'))

resized = image_ID.resize((300, 300), Image.ANTIALIAS)
photo = ImageTk.PhotoImage(resized)


### Set Frame-spacer for better overview ###
colour = 'grey'
txtcolour = 'white'
space = 2
framespacer=Frame(root1, width = space, height = 2, bg = colour)

frame_s1 = framespacer
frame_s1.grid(row=0, column=0, sticky = 'n')
frame_s2 = framespacer
frame_s2.grid(row=0, column=3, sticky = 'n')
frame_s3 = framespacer
frame_s3.grid(row=4, column=0, sticky = 'n')
frame_s4 = framespacer
frame_s4.grid(row=4, column=3, sticky = 'n')
frame_s5 = framespacer
frame_s5.grid(row=2, column=1, sticky = 'n')
frame_s6 = framespacer
frame_s6.grid(row=2, column=2, sticky = 'n')
###

frame_top = Frame(root1, width = 300, height = 300, bg = colour)
frame_top.grid(row=1, column=1)

label = Label(frame_top,image=photo)
label.image = photo
label.grid(row=0)

frame_top2 = Frame(root1, width = 300, height = 300, bg = colour)
frame_top2.grid(row=3, column=1)
frame_top2.grid_propagate(0)

frame_top3 = Frame(root1, width = 300, height = 300, bg = colour)
frame_top3.grid(row= 1, column=2, sticky = 'n')

frame_top4 = Frame(root1, width = 300, height = 300, bg = colour)
frame_top4.grid(row=3, column=2, sticky = 'n')

f = pp.figure(figsize=(2,2), dpi=150, edgecolor='k')
canvas = FigureCanvasTkAgg(f, master = frame_top3)
canvas.get_tk_widget().grid(row=0, column=0)


labtext_1 = Label(frame_top4,width = 300, height = 200,bg = colour)
labtext_1.grid(row=1, column = 0, sticky = 's')

T = Text(frame_top4, height=10, width=37)
T.grid()

#T.insert((END, "Just a text Widget\nin two lines\n"))
#s = ttk.Style()
#s.configure('TFrame', background = "#fdd57e")

#s = ttk.Style()
style = ttk.Style()
settings = {"TNotebook.Tab": 
               {"configure": {"padding": [1, 2], "background": colour , "font" : ('Sans','9','bold')}, 
                "map": {"background": [("selected", "red"),  ("active", "#fc9292")], 
                        "foreground": [("selected", "pink"), ("active", "white")], 
                        "expand": [("selected", [1, 1, 1, 1])]} },
            "TFrame": 
               {"configure": {"padding": [0, 2], "background": colour } },
            "TNotebook": 
               {"configure": {"padding": [1, 2], "background": colour } }}
                   
                   
style.theme_create('JHB', parent="alt", settings=settings)
style.theme_create('JHB2', parent="alt", settings=settings)
style.theme_use('JHB')

nb = ttk.Notebook(frame_top2, width = 7000)
nb.pressed_index = None
nb.grid(row=0, column=0, sticky='NESW')

page1 = ttk.Frame(nb)
nb.add(page1, text='Overview')

page2 = ttk.Frame(nb)
nb.add(page2, text='ROI_select')

page3 = ttk.Frame(nb)   
nb.add(page3, text='Powerseries')

page4 = ttk.Frame(nb)
nb.add(page4, text='Pinholeseries')




frame_spacer_01 = Frame(page2, width = 50, height = 1, background =colour)
frame_spacer_01.grid(row=0, column=0, sticky = 'w')

frame_5 = Frame(frame_spacer_01, width = 150, height = 110, background = colour)
frame_5.grid(row=0, column=0, sticky = 'wn')
frame_5.grid_propagate (False)

frame_6 = Frame(frame_spacer_01, width = 150, height = 180, background = colour)
frame_6.grid(row=1, column=0, sticky = 'wn')
frame_6.grid_propagate (False)

frame_7 = Frame(frame_spacer_01, width = 150, height = 180, background = colour)
frame_7.grid(row=1, column=1, sticky = 'wn')
frame_7.grid_propagate (False)

frame_8 = Frame(frame_spacer_01, width = 150, height = 110, background = colour)
frame_8.grid(row=0, column=1, sticky = 'wn')
frame_8.grid_propagate (False)


laser= Label(page3, text=' Laser:', height = 1, foreground= txtcolour, background=colour)
laser.grid(row = 0, column = 0)
laser_01 = Entry(page3,width = 8)
laser_01.insert(END, '561')
laser_01.grid(row = 0,column = 1)

laservalue= Label(page3, text=' value:', height = 1, foreground= txtcolour, background=colour)
laservalue.grid(row = 0, column = 2, sticky = 'w')
laser_value_01 = Entry(page3,width = 8)
laser_value_01.insert(END, '50')
laser_value_01.grid(row = 0, column = 3, sticky = 'w')

laser_STED= Label(page3, text=' STED-Laser:', height = 1, foreground= txtcolour, background=colour)
laser_STED.grid(row = 1, column = 0, sticky = 'w')
laser_STED_01 = Entry(page3,width = 8)
laser_STED_01.insert(END, '775')
laser_STED_01.grid(row = 1, column = 1, sticky = 'w')

laser_STEDvalue= Label(page3, text=' value:', height = 1, foreground= txtcolour, background=colour)
laser_STEDvalue.grid(row = 1, column = 2)
laser_STEDvalue_01 = Entry(page3,width = 20)
laser_STEDvalue_01.insert(END, '50, 50, 20,32')
laser_STEDvalue_01.grid(row = 1, column = 3)

#dwell= Label(frame_5, text='', height = 1 , foreground= txtcolour, background=colour)
#dwell.grid(row = 0, column = 0 , sticky = W+N)


dwell= Label(frame_5, text=' Dwelltime [us]:', height = 1 , foreground= txtcolour, background=colour)
dwell.grid(row = 0, column = 0 , sticky = W+N)
dwell_01 = Entry(frame_5,width = 3 )
dwell_01.insert(END, '1')
dwell_01.grid(row = 0, column = 1, sticky = 'w')

pxsize= Label(frame_5, text=' Pixelsize [nm]:', height = 1,foreground= txtcolour, background=colour)
pxsize.grid(row = 1, column = 0, sticky = W+N)
pxsize_01 = Entry(frame_5, width = 3)
pxsize_01.insert(END, '10')
pxsize_01.grid(row = 1, column = 1, sticky = 'w')

ROIsize= Label(frame_5, text=' ROIsize [um]:', height = 1,foreground= txtcolour, background=colour)
ROIsize.grid(row = 2, column = 0, sticky = W+N)
ROIsize_01 = Entry(frame_5,width = 3)
ROIsize_01.insert(END, '1')
ROIsize_01.grid(row = 2, column = 1, sticky = 'w')

frames= Label(frame_5, text=' # Frames:           ', height = 1, foreground= txtcolour, background=colour)
frames.grid(row = 3, column = 0, sticky = W+N)
frames_01 = Entry(frame_5,width = 3)
frames_01.insert(END, '21')
frames_01.grid(row = 3, column = 1, sticky = 'w')


MultiRUN= Label(frame_8, text=' Focus search:', height = 1, foreground= txtcolour, background=colour)
MultiRUN.grid(row = 0, column = 2, sticky = 'w')
MultiRUN_01 = Entry(frame_8,width = 8)
MultiRUN_01.insert(END, '3')
MultiRUN_01.grid(row = 0, column = 3, sticky = 'w')

MultiRUN_s= Label(frame_8, text='', height = 1, foreground= txtcolour, background=colour)
MultiRUN_s.grid(row = 0, column = 0, sticky = 'w')

Autofocus= Label(frame_8, text=' Autofocus:', height = 1 ,foreground= txtcolour, background=colour)
Autofocus.grid(row = 2, column = 2, sticky = W)
var13 = IntVar()
Autofocus_01 = Checkbutton(frame_8,  variable = var13, background =colour)
Autofocus_01.grid(row=2, column = 3, sticky = W)

QFS= Label(page1, text=' Autofocus:', height = 1 ,foreground= txtcolour, background=colour)
QFS.grid(row = 2, column = 2, sticky = W)
var14 = IntVar()
QFS_01 = Checkbutton(page1,  variable = var14, background =colour)
QFS_01.grid(row=2, column = 3, sticky = W)


L485= Label(frame_6, text=' L485:', height = 1 ,foreground= txtcolour, background=colour)
L485.grid(row = 0, column = 0, sticky = W)
var1 = IntVar()
L485_01 = Checkbutton(frame_6,  variable = var1, background =colour)
L485_01.grid(row=0, column = 1, sticky = W)

L485_02= Label(frame_6, text=' L485:', height = 1 ,foreground= txtcolour, background=colour)
L485_02.grid(row = 0, column = 0, sticky = W)
var7 = IntVar()
L485_02_ = Checkbutton(frame_6,  variable = var7, background =colour)
L485_02_.grid(row=0, column = 2, sticky=W)

L518= Label(frame_6, text=' L518:', height = 1 ,foreground= txtcolour, background=colour)
L518.grid(row = 1, column = 0, sticky = W+N)
var2 = IntVar()
L518_01 = Checkbutton(frame_6,  variable = var2, background =colour)
L518_01.grid(row=1, column = 1, sticky=W)

L518_02= Label(frame_6, text=' L518:', height = 1,foreground= txtcolour, background=colour)
L518_02.grid(row = 1, column = 0, sticky = W+N)
var8 = IntVar()
L518_02_ = Checkbutton(frame_6,  variable = var8, background =colour)
L518_02_.grid(row=1, column = 2, sticky=W)

L561= Label(frame_6, text=' L561:', height = 1,foreground= txtcolour, background=colour)
L561.grid(row = 2, column = 0, sticky = W+N)
var3 = IntVar()
L561_01 = Checkbutton(frame_6,  variable = var3, background =colour)
L561_01.grid(row=2, column = 1, sticky=W)

L561_02= Label(frame_6, text=' L561:', height = 1,foreground= txtcolour, background=colour)
L561_02.grid(row = 2, column = 0, sticky = W+N)
var9 = IntVar()
L561_02_ = Checkbutton(frame_6,  variable = var9, background =colour)
L561_02_.grid(row=2, column = 2, sticky=W)

L640= Label(frame_6, text=' L640:', height = 1,foreground= txtcolour, background=colour)
L640.grid(row = 3, column = 0, sticky = W+N)
var4 = IntVar()
L640_01 = Checkbutton(frame_6, variable = var4, background =colour)
L640_01.grid(row=3, column = 1, sticky=W)

L640_02= Label(frame_6, text=' L640:', height = 1,foreground= txtcolour, background=colour)
L640_02.grid(row = 3, column = 0, sticky = W+N)
var10 = IntVar()
L640_02_ = Checkbutton(frame_6,  variable = var10, background =colour)
L640_02_.grid(row=3, column = 2, sticky = W+N)

L595= Label(frame_6, text=' L595:', height = 1,foreground= txtcolour, background=colour)
L595.grid(row = 4, column = 0, sticky = W+N)
var5 = IntVar()
L595_01 = Checkbutton(frame_6,  variable = var5, background =colour)
L595_01.grid(row=4, column = 1, sticky = W+N)

L595_02= Label(frame_6, text=' L595:', height = 1,foreground= txtcolour, background=colour)
L595_02.grid(row = 4, column = 0, sticky = W+N)
var11 = IntVar()
L595_02_ = Checkbutton(frame_6,  variable = var11, background =colour)
L595_02_.grid(row=4, column = 2, sticky = W+N)

L775= Label(frame_6, text=' L775:', height = 1,foreground= txtcolour, background=colour)
L775.grid(row = 5, column = 0, sticky = W+N)
var6 = IntVar()
L775_01 = Checkbutton(frame_6, variable = var6, background =colour)
L775_01.grid(row=5, column = 1, sticky = W+N)

L775_02= Label(frame_6, text=' L775:', height = 1,foreground= txtcolour, background=colour)
L775_02.grid(row = 5, column = 0, sticky = W+N)
var12 = IntVar()
L775_02_ = Checkbutton(frame_6,  variable = var12, background =colour)
L775_02_.grid(row=5, column = 2, sticky = W+N)

L485_value= Label(frame_6, height = 1,foreground= txtcolour, background=colour)
L485_value.grid(row = 2, column = 0, sticky = W+N)
L485_value_01 = Entry(frame_6,width = 3)
L485_value_01.insert(END, '0')
L485_value_01.grid(row = 0, column = 5, sticky = E+N)

L518_value= Label(frame_6, height = 1,foreground= txtcolour, background=colour)
L518_value.grid(row = 1, column = 0, sticky = W+N)
L518_value_01 = Entry(frame_6,width = 3)
L518_value_01.insert(END, '0')
L518_value_01.grid(row = 1, column = 5, sticky = E+N)

L561_value= Label(frame_6, height = 1,foreground= txtcolour, background=colour)
L561_value.grid(row = 2, column = 0, sticky = W+N)
L561_value_01 = Entry(frame_6,width = 3)
L561_value_01.insert(END, '0')
L561_value_01.grid(row = 2, column = 5, sticky = E+N)

L640_value= Label(frame_6, height = 1,foreground= 'black', background=colour)
L640_value.grid(row = 3, column = 0, sticky = W+N)
L640_value_01 = Entry(frame_6,width = 3)
L640_value_01.insert(END, '0')
L640_value_01.grid(row = 3, column = 5, sticky = E+N)

L595_value= Label(frame_6, height = 1,foreground= 'black', background= colour)
L595_value.grid(row = 4, column = 0, sticky = W+N)
L595_value_01 = Entry(frame_6,width = 3)
L595_value_01.insert(END, '0')
L595_value_01.grid(row = 4, column = 5, sticky = E+N)

L775_value= Label(frame_6, height = 1,foreground= 'black', background=colour)
L775_value.grid(row = 5, column = 0, sticky = W+N)
L775_value_01 = Entry(frame_6,width = 3)
L775_value_01.insert(END, '0')
L775_value_01.grid(row = 5, column = 5, sticky = E+N)

THRES_value= Label(frame_6, height = 1,foreground= 'black', background=colour)
THRES_value.grid(row = 6, column = 0, sticky = W+N)
THRES_value_01 = Entry(frame_6,width = 3)
THRES_value_01.insert(END, '0')
THRES_value_01.grid(row = 6, column = 5, sticky = E+N)

a=0 


def SSS(scale_03_value):
    Findpeak()
    
def TEST():
    
    global pixelsize
    global Roisize
    global dwelltime
    global frame_number
    
    global act485
    global act518
    global act561
    global act640
    global act595
    global act775
    
    global act485_02
    global act518_02
    global act561_02
    global act640_02
    global act595_02
    global act775_02
    global act_Autofocus
    global act_QFS
    
    
    pixelsize = pxsize_01.get()
    Roisize= ROIsize_01.get()
    dwelltime = dwell_01.get()
    frame_number = frames_01.get()
    
    act485 = var1.get()
    act518 = var2.get()
    act561 = var3.get()
    act640 = var4.get()
    act595 = var5.get()
    act775 = var6.get()
    
    act485_02 = var7.get()
    act518_02 = var8.get()
    act561_02 = var9.get()
    act640_02 = var10.get()
    act595_02 = var11.get()
    act775_02 = var12.get()
    act_Autofocus = var13.get()
    act_QFS = var14.get()
    
    
 
    
def Connect():
    global a
    a=0
    try:
        im=specpy.Imspector()
    except (RuntimeError):
        a=1
        
    if a==0:
        T.insert(END, 'Connected')
    else:
        T.insert(END, 'Disconnected')
    
    
def Findpeak():
    global x_new
    global y_new
    global R
    global scale_val
    global aa
    global number_peaks_new 
    global time_wait_Multirun
    
    global x_transfer
    global y_transfer
    global CO
    
    data = Image.open('{}{}'.format('D:/current data/','Overview_image.tiff'))
    
    data_sz = data.size
    zz = np.zeros(data_sz)
    aa = roi_size
    
    #threshold = float(THRES_value_01.get())
    threshold = scale_01.get()      # define the lower threshold value
    print("threshold value is %.1f" % threshold)
    R_min = scale_02.get()          # additional distance threshold
    R_max = scale_03.get()
    exc_border_peaks = 1 # type 0 for activate for border peaks 
############################################
    #Calculate parameters    
############################################

    if os.path.exists('{}{}'.format(path,foldername)) == False:
        os.makedirs('{}{}'.format(path,foldername))
        os.makedirs('{}{}'.format(path, 'test_images.tiff'))
        

    imarray = np.array(data)
    imarray_thres = stats.threshold(imarray, threshmin= threshold, newval=0)
    array_size_x = imarray_thres.shape[1]
    array_size_y = imarray_thres.shape[0]
    
    v1_x = int(array_size_x/2)
    v1_y = int(array_size_x/2)
    
    ##############################################
    # Find peaks in images:   
    coordinates = peak_local_max(imarray_thres, min_distance= R_min, indices=True) #exclude_border=exc_border_peaks)
    y_coord = coordinates[:, 0]
    x_coord = coordinates[:, 1]
     
    if len(coordinates)<= 1:
        CO=1
        print('NearestNeighbors failed: coordinates is <= 1')
        x_transfer = []
        y_transfer = []
        np.savetxt('{}{}{}'.format('D:/current data/','x_transfer','.dat'), x_transfer)       
        np.savetxt('{}{}{}'.format('D:/current data/','y_transfer','.dat'), y_transfer) 
        
    else:
        CO= 0
        print('NearestNeighbors okay')
        nbrs = NearestNeighbors(n_neighbors=3, algorithm='ball_tree').fit(coordinates)
        distances, indices = nbrs.kneighbors(coordinates) 
            
        number_peaks = int(coordinates.size/ 2)
        t_x= np.zeros(number_peaks)
        t_y= np.zeros(number_peaks)
        R_new = np.zeros(number_peaks)
        
        for i in range(number_peaks):
            for ii in range(number_peaks):
                R = math.sqrt((coordinates[i][0]-coordinates[ii][0])**2+(coordinates[i][1]-coordinates[ii][1])**2)
                
                if R < R_max and R > 0:#R_min:
                    #print ('peaks',i,ii,'x1_coord',coordinates[i][0],'y1_coord',coordinates[i][1],'x2_coord',coordinates[ii][0],'y2_coord',coordinates[ii][1],'Distance R',R,'below', R_min  )
                    #print (coordinates[i][0],coordinates[i][1])
                    t_y[i] = coordinates[i][0]
                    t_x[i] = coordinates[i][1]
                    R_new[i] = R
  
        R_new_thres = R_new[np.where(R_new>0)]
    
        ###################################
        # Libary import
            
        # Do you want to simulate a Poisson distribution of spots [own_data = 0] or use own data [own_data = 1]????
        own_data = 1
             
        ###############  Window parameters  ##########################################
        pixelsize= 1 #[nm]
        
        xMin=0
        xMax=data.size[0]; # pixelnumber in x
        yMin=0
        yMax=data.size[1]; # pixelnumber in y
         
        #Poisson process parameters simulation
        lambda0=5; #intensity (ie mean density) of the Poisson process
         
        ###############################################################################
        if own_data ==0:
            print('Poisson distribution of spots')
            #Simulate Poisson point process
            xDelta=xMax-xMin;yDelta=yMax-yMin; #rectangle dimensions
            areaTotal=xDelta*yDelta; #total area
            numbPoints = scipy.stats.poisson( lambda0*areaTotal ).rvs()#Poisson number of points
            x_Poisson = xDelta*scipy.stats.uniform.rvs(0,1,((numbPoints,1)))+xMin#x coordinates of Poisson points
            y_Poisson = yDelta*scipy.stats.uniform.rvs(0,1,((numbPoints,1)))+yMin#y coordinates of Poisson points
            xx = x_Poisson
            yy = y_Poisson
            
        else:
            print('Using real data for statistics')
            xDelta=xMax-xMin;yDelta=yMax-yMin; #rectangle dimensions
            areaTotal=xDelta*yDelta; #total area
                
       
        z = coordinates
        #np.savetxt('C:/Users/buddeja/Desktop/test/z_GUI.txt', z, delimiter='\t')   # X is an array
        #print(z.shape)
        Kest = RipleysKEstimator(area=areaTotal, x_max=xMax, y_max=yMax, x_min=xMin, y_min=yMin)
        r = np.linspace(0, xMax, 1000)
        ###########################################################################
    
    
        f1 = pp.figure(figsize=(3,3), dpi=100, edgecolor='k',facecolor = 'r')
        LABELSIZE = 7
        a = f1.add_subplot(211)
        a.hist(distances[:,1], bins=20,color='blue', label='1st neighbor',alpha=0.2)
        a.hist(distances[:,2], bins=20,color='red', label='2nd neighbor',alpha=0.2)
        pp.xlabel('Nearst neighbor distance [a.u]', fontsize = LABELSIZE, fontweight='bold')
        pp.ylabel('counts', fontsize = LABELSIZE, fontweight='bold')
        a.xaxis.set_tick_params(labelsize=LABELSIZE)
        a.yaxis.set_tick_params(labelsize=LABELSIZE)
        pp.legend(fontsize = LABELSIZE)
    
        
        y = r*0
        b = f1.add_subplot(212)
        #b.plot(r*pixelsize, Kest(data=z, radii=r, mode='translation'), color='black',label=r'$K_{trans}$')
        b.plot(r*pixelsize, Kest.Hfunction(data=z, radii=r, mode='translation'), color='black',label=r'$K_{trans}$')
        b.plot(r*pixelsize, y, color='red', ls='--', label=r'$K_{zerosline}$')
        #b.plot(r*pixelsize, Kest.poisson(r), color='green', ls=':', label=r'$K_{pois}$')
        pp.xlabel('radius [a.u]', fontsize = LABELSIZE, fontweight='bold')
        pp.ylabel('H-function', fontsize = LABELSIZE, fontweight='bold')
        pp.text(5, 70, r'Clustering')
        pp.text(5, -90, r'Dispersion')
        #pp.text(500, 70, '{}{}'.format('\u03C1=',x_pixelsize))
        b.set_xlim([0,999])
        b.set_ylim([-100,100])
        b.xaxis.set_tick_params(labelsize=LABELSIZE)
        b.yaxis.set_tick_params(labelsize=LABELSIZE)
        pp.legend(fontsize = LABELSIZE)
    
        pp.subplots_adjust(bottom=0.12, right=0.98, left=0.2,  top=0.95, wspace= 0.5, hspace= 0.43)
        
        canvas = FigureCanvasTkAgg(f1, master = frame_top3)
        canvas.get_tk_widget().grid(row=0, column=0, rowspan=1)
     
    
        x_new = t_x[np.where(t_x > 0 )]        
        y_new = t_y[np.where(t_y > 0 )]
        
        s = np.zeros(number_peaks)
        x = np.zeros(number_peaks)
        y = np.zeros(number_peaks)
        
        for i in range(R_new_thres.size):
            if R_new_thres[i] not in s:
                s[i]= R_new_thres[i]
                x[i]= x_new[i]
                y[i]= y_new[i]
            else:
                s[i]= 0
                x[i]= 0
                y[i]= 0
                 
        s = s[np.where(s > 0 )] 
         
              
    
        x_new = x[np.where(x > 0 )]
        y_new = y[np.where(y > 0 )]
        
        x_transfer = x_new - v1_x 
        y_transfer = y_new - v1_y
        
        number_peaks_new = x_new.size
    
        ##########################################
        # printing and saving the results
        
        np.savetxt('{}{}{}'.format('D:/current data/','x_coord','.dat'), x_coord)      
        np.savetxt('{}{}{}'.format('D:/current data/','y_coord','.dat'), y_coord)   
        
        np.savetxt('{}{}{}'.format('D:/current data/','x_new','.dat'), x_new)         
        np.savetxt('{}{}{}'.format('D:/current data/','y_new','.dat'), y_new)
        
        np.savetxt('{}{}{}'.format('D:/current data/','x_transfer','.dat'), x_transfer)       
        np.savetxt('{}{}{}'.format('D:/current data/','y_transfer','.dat'), y_transfer) 
    

    
        for i in range(x_new.size):
            zz[int(y_new[i]),int(x_new[i])] = 1
    
        sigma_1 = 4
        sigma = [sigma_1, sigma_1]   
        y = gaussian_filter(zz, sigma, mode='constant')
        yy= stats.threshold(y, threshmax=0, newval=255)
        #pp.imshow(yy, 'hot')
    
        xxx = Image.fromarray(yy)#imarray_thres)
        resized = xxx.resize((300, 300), Image.ANTIALIAS)
        
        
        #########
        
        
        #img = Image.open('{}{}{}'.format('C:/Users/Abberior_admin/Desktop/GUI_20192611/GUI_pythoninterface/','Overview_image','.tiff'))
        img = Image.open('{}{}'.format('D:/current data/','Overview_image.tiff'))
        im = np.array(img)
        im = cm.hot(im)
        im = np.uint8(im * 255)
        t = Image.fromarray(im)
        A = t.convert("RGBA")
        
        im = yy
        im = cm.Greens(im)
        im = np.uint8(im * 255)
        t = Image.fromarray(im)
        B = t.convert("RGBA")
        
        C = Image.blend(A, B, alpha=.3)
        C = C.resize((300, 300), Image.ANTIALIAS)
        
        photo = ImageTk.PhotoImage(C)


#########################
    #photo = ImageTk.PhotoImage(resized)
  
    
        label = Label(frame_top,image=photo)
        label.image = photo
        label.grid(row=0)
        
        AREA_spots = 0.09
        rho= number_peaks_new/((roi_size*1000000)*(roi_size*1000000))
        Lbda= rho*AREA_spots
        
        TEST = np.zeros((4,4))
        
        
        for i in range(4):
            P= (Lbda**i)/math.factorial(i)*math.exp(- Lbda )
            TEST[0,i]=i
            TEST[1,i]=Lbda
            TEST[2,i]=P
            TEST[3,i]=i
        
        print('Prop. of Singlet-events:',np.around(100-(TEST[2,1]/np.sum(TEST[2,1:])*100), decimals=2),'%')  
            
            
        
        
        T.delete('1.0', END)  
        T.insert(END,'{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}{}'.format('peaks_found:  ',number_peaks_new-1,'\n',
                                                             '\n',
                                                             'Surface-density [µm^(-2)]:',rho,'\n',
                                                             'Area of spots [µm^(2)]:',AREA_spots,'\n',
                                                             'Total Area [µm^(2)]:',((roi_size*1000000)*(roi_size*1000000)),'\n',
                                                             'Prop. of Singlet-events: ',np.around((TEST[2,1]/np.sum(TEST[2,1:])*100), decimals=2),'%','\n',
                                                             'Prop. of Multi-events:   ',np.around(100-(TEST[2,1]/np.sum(TEST[2,1:])*100), decimals=2),'%'))
        
        time_wait_Multirun = number_peaks_new

      
    
def Overview(Multi, Pos):
    
    import time
    import numpy as np
    import os
    import time
    import PIL.Image as Image
    import matplotlib.pyplot as pp
    import specpy
    global x_pixelsize
    global roi_size
    
    if os.path.exists('{}{}'.format(path,foldername)) == False:
        os.makedirs('{}{}'.format(path,foldername))    
    else:
        shutil.rmtree('{}{}{}'.format(path,foldername,'/'))#(D:/current data/testfolder')
        os.makedirs('{}{}'.format(path,foldername))
        
        
    # Define all  measurement relevant paramentetrs

    # Relative ROI coordinates from  find peak in tiff
    z_position = 0        # in meter
    x_pixelsize = 5e-08      # in meter
    y_pixelsize = 5e-08       # in meter
    z_pixelsize = 5e-08       # in meter
    roi_size = 20e-06          # in meter
    
    Dwelltime= 20e-06         # in seconds  
    LP561 = 0
    LP640 = 20
    
    Streaming_HydraHarp= True # Type True or False for Streaming via HydraHarp
    
    modelinesteps= True      #Type True for linesteps
    number_linesteps = 2     #Type the number of linesteps
    xyt_mode = 1296           #Type 785 for xyt mode  784 
    
    number_frames = 1      #Type number of frames t in xyt mode
    
    time_wait = math.ceil((roi_size/x_pixelsize) * (roi_size/y_pixelsize)* number_frames* Dwelltime) + 1
    
    Activate485 = False                    
    Activate518 = False  
    Activate561 = True  
    Activate640 = True  
    Activate595 = False  
    Activate775 = True 
    #Activate_QFS = bool(act_QFS)
    
    if a==0:
        
            
        im = specpy.Imspector()    
        c=im.create_measurement()
                
        if Multi == 1:
            print('MULTIRUN')
            c.set_parameters('ExpControl/scan/range/offsets/coarse/y/g_off', Pos) # current stage position in x [m]
        else:
            print('non MULTIRUN')
            
        #############    
#        if Activate_QFS == True:
#            M_obj= im.measurement(c.name())
            
        #########################    
        c.set_parameters('ExpControl/scan/range/mode',xyt_mode)
        c.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', True)
        c.set_parameters('ExpControl/scan/range/x/psz',x_pixelsize)
        c.set_parameters('ExpControl/scan/range/y/psz',y_pixelsize)
        c.set_parameters('ExpControl/scan/range/z/psz',z_pixelsize)
        c.set_parameters('ExpControl/scan/range/z/off', 1e-15)#z_position*z_pixelsize)
        c.set_parameters('ExpControl/scan/range/x/len',roi_size)
        c.set_parameters('ExpControl/scan/range/y/len',roi_size)
        c.set_parameters('ExpControl/scan/range/z/len',roi_size)
        c.set_parameters('ExpControl/scan/dwelltime', Dwelltime)
        c.set_parameters('ExpControl/scan/range/t/res',number_frames)
        c.set_parameters('HydraHarp/data/streaming/enable', False)
        
        c.set_parameters('ExpControl/lasers/power_calibrated/3/value/calibrated', LP561)
        c.set_parameters('ExpControl/lasers/power_calibrated/4/value/calibrated', LP640)
        c.set_parameters('ExpControl/gating/linesteps/laser_enabled',[False, False, False, False, True, False, False, False])
        c.set_parameters('ExpControl/gating/linesteps/laser_on',[[False, False, False, False, True, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False],
         [False, False, False, False, False, False, False, False]])
        
        
        c.set_parameters('ExpControl/gating/linesteps/chans_enabled',[True,False,False,False])
        c.set_parameters('ExpControl/gating/linesteps/step_values',[1,0,0,0,0,0,0,0])
        
        c.set_parameters('HydraHarp/is_active', True)
        c.set_parameters('ExpControl/gating/tcspc/channels/0/mode', 0)
        c.set_parameters('ExpControl/gating/tcspc/channels/0/stream', 'HydraHarp')
        c.set_parameters('ExpControl/gating/tcspc/channels/1/mode', 0)
        c.set_parameters('ExpControl/gating/tcspc/channels/1/stream', 'HydraHarp')
        c.set_parameters('ExpControl/gating/tcspc/channels/2/mode', 0)
        c.set_parameters('ExpControl/gating/tcspc/channels/2/stream', 'HydraHarp')
        c.set_parameters('ExpControl/gating/tcspc/channels/3/mode', 0)
        c.set_parameters('ExpControl/gating/tcspc/channels/3/stream', 'HydraHarp')
        c.set_parameters('ExpControl/scan/detsel/detsel',['APD4', 'APD2', 'APD4', 'APD1'])
    #    M_name= im.measurement_names()[0]
    #    M_obj= im.measurement(M_name)
    #    im.run(M_obj)
        meas= im.active_measurement()
        im.run(meas)
        
        time.sleep(time_wait)
        meas = im.active_measurement()
        stk_names = meas.stack_names()
        stk = meas.stack(stk_names[0])
        
        pix_data = stk.data()
        #mean = pix_data.mean()
        print(pix_data.shape)
        if pix_data.shape == (1,1,400,400):
            pix_data = gaussian_filter(pix_data, 1)
            io = Image.fromarray(pix_data[0,0]*10)
            print(max(pix_data.flatten()))
            pp.imshow(io)
            data = io
        elif pix_data.shape == (1, 400, 400, 1):
            pix_data = gaussian_filter(pix_data, 1)
            io = Image.fromarray(pix_data[0,:,:,0]*10)
            print(max(pix_data.flatten()))
            pp.imshow(io)
            data = io
        else:
            io = Image.fromarray((np.mean(pix_data, axis=3)[0])*10) # create greyscale image object  

            pp.imshow(io)
            data = io
        
    else:
        data = Image.open('{}{}'.format(path,'TEST_01.tiff'))
        #print('datasize = ', data.size)
        io_1 = np.array(data)
        io = Image.fromarray(io_1 ) 
        
        if Multi == 1:
            print('MULTIRUN')
            #print('ExpControl/scan/range/offsets/coarse/y/g_off', Pos) # current stage position in x [m]
        else:
            print('non MULTIRUN')
        
    
    scale_01.config(to =np.max(io))            
    io.save('{}{}'.format('D:/current data/','Overview_image.tiff'), format = 'TIFF')#'D:/current data/','Overview_image.tiff'
    
    im = np.array(data)
    im = cm.hot(im)
    im = np.uint8(im * 255)
    t = Image.fromarray(im)
    resized = t.resize((300, 300), Image.ANTIALIAS)
    
    photo = ImageTk.PhotoImage(resized)
    
    
    label = Label(frame_top,image=photo)
    label.image = photo
    label.grid(row=0)
    
    T.delete('1.0', END) 
    T.insert(END, "Overview created")

def powerseries():
    im = specpy.Imspector()
    import time
    
    powervector = int(laser_value_01.get())
    powervector_STED = [int(s) for s in laser_STEDvalue_01.get().split(',')]
    laser = int(laser_01.get()) 
    laser_STED = int(laser_STED_01.get())
    
    
    if laser == 485:laservalue = 0
    elif laser == 518:laservalue = 1
    elif laser == 595:laservalue = 2
    elif laser == 561:laservalue = 3
    elif laser == 640:laservalue = 4
    elif laser == 775:laservalue = 5
    
    if laser_STED == 485:laser_STEDvalue = 0
    elif laser_STED == 518:laser_STEDvalue = 1
    elif laser_STED == 595:laser_STEDvalue = 2
    elif laser_STED == 561:laser_STEDvalue = 3
    elif laser_STED == 640:laser_STEDvalue = 4
    elif laser_STED == 775:laser_STEDvalue = 5
  
          
    for ii in powervector_STED:
        c=im.create_measurement()
        #print(ii)
        c.set_parameters('{}{}{}'.format('ExpControl/lasers/power_calibrated/', laservalue,'/value/calibrated'), powervector)
        c.set_parameters('{}{}{}'.format('ExpControl/lasers/power_calibrated/', laser_STEDvalue,'/value/calibrated'), ii)
        M_name= im.measurement_names()[ii + 1]
        M_obj= im.measurement(M_name)
        im.run(M_obj)
        
        time.sleep(2) # timewait set to 5 sec
 
       
def pinholeseries():
    im = specpy.Imspector()
    
    pinholevector = [1,2,3,4,50]
    laser = 3
    
    for ii in pinholevector:
        c=im.create_measurement()
        c.set_parameters('{}{}{}'.format('ExpControl/lasers/pinholeseries'), ii)
        M_name= im.measurement_names()[i + 1]
        M_obj= im.measurement(M_name)
        im.run(M_obj)
        time.sleep(5) # timewait set to 5 sec
    
 

def Run_meas(Multi, Pos): 
    import math
    import time
    global save_path
    
 # Define all  measurement relevant paramentetrs
    im = specpy.Imspector() 
        
 # Relative ROI coordinates from  find peak in tiff
    z_position =  5e-08 #float(pixelsize)*1e-09         # in meter
    x_pixelsize = float(pixelsize)*1e-09        # in meter
    y_pixelsize = float(pixelsize)*1e-09        # in meter
    z_pixelsize = float(pixelsize)*1e-09        # in meter
    roi_size =    float(Roisize)*1e-06            # in meter 
    Dwelltime= float(dwelltime)*1e-06         # in seconds  
        
    LP485 = int(L485_value_01.get())
    LP518 = int(L518_value_01.get())
    LP561 = int(L561_value_01.get())
    LP640 = int(L640_value_01.get())
    LP595 = int(L595_value_01.get())
    LP775 = int(L775_value_01.get())
            
    Streaming_HydraHarp= True # Type True or False for Streaming via HydraHarp 
    modelinesteps= True      #Type True for linesteps
    number_linesteps = 2     #Type the number of linesteps
    xyt_mode = 784           #Type 785 for xyt mode
        
    number_frames = float(frame_number)     #Type number of frames t in xyt mode
    time_wait = math.ceil((roi_size/x_pixelsize) * (roi_size/y_pixelsize)* number_frames* Dwelltime) + 1
                
    Activate485 = bool(act485)                    
    Activate518 = bool(act518)   
    Activate561 = bool(act561)   
    Activate640 = bool(act640)   
    Activate595 = bool(act595)   
    Activate775 = bool(act775)   
        
    Activate485_02 = bool(act485_02)                    
    Activate518_02 = bool(act518_02)   
    Activate561_02 = bool(act561_02)   
    Activate640_02 = bool(act640_02)   
    Activate595_02 = bool(act595_02)   
    Activate775_02 = bool(act775_02) 
    #Activate_Autofocus= bool(act_Autofocus)
    
    
        
    ############################
    #Hydraharp-setting
        
    if Streaming_HydraHarp == True:
        stream = 'HydraHarp'
    else:
        stream = 'fpga'
                   
    Ch1_mode= 0                # type 0 for counter and 1 for flim
    Ch1_stream= stream         # type hydraharp or fpga for internal
    Ch2_mode= 0          
    Ch2_stream= stream
    Ch3_mode= 0         
    Ch3_stream= stream
    Ch4_mode= 0         
    Ch4_stream= stream
        
    ############################
    #detector activation
    
    detector1= True
    detector2= True
    detector3= True
    detector4= True
        
   ####################################################################################
        
    T.delete('1.0', END) 
    T.insert(END, 'L485',L485_value_01.get())
        
    
    linesteps = [False, False, False, False, False, False, False, False]
    for i in range(number_linesteps):
        linesteps[i] = True
        
    #M_obj= im.measurement(im.measurement_names()[0])
    
    d = im.active_measurement()
    M_obj= im.measurement(d.name())
             
#    if Activate_Autofocus == True: #Activate_Autofocus
#        M_obj.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', False)
#        time.sleep(2) 
#        M_obj.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', True)
        
    
            
                             
    if CO ==1 :
        print('failed')
        save_path = '{}{}{}{}{}'.format('D:/current data/','Overview_',Pos,'_numberSPOTS_', 0)
        os.makedirs(save_path)  
        (im.measurement(im.measurement_names()[1])).save_as('{}{}{}{}{}'.format(save_path,'/','Overview_', Pos,'.msr'))
        im.close(im.measurement(im.measurement_names()[1]))
    else:
        x_global_offset = M_obj.parameters('ExpControl/scan/range/x/off')
        y_global_offset = M_obj.parameters('ExpControl/scan/range/y/off')
        x_roi_new = np.loadtxt('D:/current data/x_transfer.dat') - x_global_offset
        y_roi_new = np.loadtxt('D:/current data/y_transfer.dat') - y_global_offset
        for i in range(1,x_roi_new.size):
            x_position = x_roi_new[i]
            y_position = y_roi_new[i]
               
            d = im.active_measurement()
            M_obj.clone(d.active_configuration())
            M_obj.activate(M_obj.configuration(i))          
            c = M_obj.configuration(i)
            
            c.set_parameters('ExpControl/scan/range/x/psz',x_pixelsize)
            c.set_parameters('ExpControl/scan/range/y/psz',y_pixelsize)
            c.set_parameters('ExpControl/scan/range/z/psz',z_pixelsize)
            c.set_parameters('ExpControl/scan/range/x/off', x_position*5e-08)
            c.set_parameters('ExpControl/scan/range/y/off', y_position*5e-08)
            c.set_parameters('ExpControl/scan/range/z/off', 1e-15)#z_position*z_pixelsize)
            c.set_parameters('ExpControl/scan/range/x/len',roi_size)
            c.set_parameters('ExpControl/scan/range/y/len',roi_size)
            c.set_parameters('ExpControl/scan/range/z/len',roi_size)
            c.set_parameters('ExpControl/scan/dwelltime', Dwelltime)
            c.set_parameters('HydraHarp/data/streaming/enable', Streaming_HydraHarp)
            c.set_parameters('HydraHarp/is_active', Streaming_HydraHarp)
            c.set_parameters('ExpControl/gating/tcspc/channels/0/mode', Ch1_mode)
            c.set_parameters('ExpControl/gating/tcspc/channels/0/stream', Ch1_stream)
            c.set_parameters('ExpControl/gating/tcspc/channels/1/mode', Ch2_mode)
            c.set_parameters('ExpControl/gating/tcspc/channels/1/stream', Ch2_stream)
            c.set_parameters('ExpControl/gating/tcspc/channels/2/mode', Ch3_mode)
            c.set_parameters('ExpControl/gating/tcspc/channels/2/stream', Ch3_stream)
            c.set_parameters('ExpControl/gating/tcspc/channels/3/mode', Ch4_mode)
            c.set_parameters('ExpControl/gating/tcspc/channels/3/stream', Ch4_stream)
            c.set_parameters('ExpControl/gating/linesteps/on', modelinesteps)
            c.set_parameters('ExpControl/gating/linesteps/steps_active', linesteps)
            c.set_parameters('ExpControl/gating/linesteps/step_values',[1,1,0,0,0,0,0,0])
            c.set_parameters('ExpControl/scan/range/mode',xyt_mode)
            c.set_parameters('ExpControl/lasers/power_calibrated/3/value/calibrated', LP561)
            c.set_parameters('ExpControl/lasers/power_calibrated/4/value/calibrated', LP640)
            c.set_parameters('ExpControl/lasers/power_calibrated/5/value/calibrated', LP775)
            c.set_parameters('ExpControl/gating/pulses/pulse_chan/delay',[0.0, 0.0, 0.0, 0.0])
            #c.set_parameters('ExpControl/lasers/power_calibrated/5/value/calibrated', LP775)
            c.set_parameters('ExpControl/gating/linesteps/laser_enabled',[True, True, True, True, True, True, False, False])
            c.set_parameters('ExpControl/gating/linesteps/laser_on',[[Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485_02, Activate518_02, Activate595_02, Activate561_02, Activate640_02, Activate775_02, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
             [Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False]])
            c.set_parameters('ExpControl/scan/range/t/res',number_frames)
            c.set_parameters('ExpControl/gating/linesteps/chans_enabled',[detector1,detector2,detector3,detector4])
            c.set_parameters('ExpControl/scan/detsel/detsel',['APD1', 'APD2', 'APD3', 'APD4'])
            c.set_parameters('ExpControl/gating/linesteps/chans_on', [[True, True, True, True],
             [True, True, True, True],
             [False, False, False, False],
             [False, False, False, False],
             [False, False, False, False],
             [False, False, False, False],
             [False, False, False, False],
             [False, False, False, False]])
        
    
              #### # c.set_parameters('{}{}{}'.format('ExpControl/lasers/pinholeseries'), ii)
               
    
            im.run(M_obj)
            time.sleep(time_wait) 
                
    #        
        if Multi ==0:
            save_path = '{}{}{}'.format('D:/current data/','Overview','.msr')
            (im.measurement(im.measurement_names()[0])).save_as(save_path)
            im.close(im.measurement(im.measurement_names()[0])) 
            files = os.listdir('D:/current data/')
            files_txt = [i for i in files if i.endswith('.ptu')]
            n_files_txt = len(files_txt)
                
            for ii in range(n_files_txt):
                new_path = '{}{}{}{}'.format('D:/current data/testfolder/ptu_files/','measurement_',ii,'.ptu')
                os.rename('{}{}'.format('D:/current data/',files_txt[ii]), '{}{}{}{}{}'.format('D:/current data/','Overview_','spot_', ii, '.ptu'))
                
        else:
            save_path = '{}{}{}{}{}'.format('D:/current data/','Overview_',Pos,'_numberSPOTS_', number_peaks_new-1)
            os.makedirs(save_path)  
            (im.measurement(im.measurement_names()[1])).save_as('{}{}{}{}{}'.format(save_path,'/','Overview_', Pos,'.msr'))
             
            files = os.listdir('D:/current data/')
            files_ptu = [i for i in files if i.endswith('.ptu')]
            files_dat = [i for i in files if i.endswith('.dat')]
            
            n_files_ptu = len(files_ptu)
            n_files_dat = len(files_dat)
                
            for ii in range(n_files_ptu):
                os.rename('{}{}'.format('D:/current data/',files_ptu[ii]), '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, '_spot_', ii, '.ptu'))
            
            for ii in range(n_files_dat):
                os.rename('{}{}'.format('D:/current data/',files_dat[ii]), '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, '_spot_', files_dat[ii],'.dat'))
        
        #np.savetxt('D:/current data/Settings.txt', [c.parameters('ExpControl')], fmt='%s')
       
def SAVING(path):

#    import PIL
#    
#    from specpy import *
#    import specpy
    
    #### functions #####
    def GT(m_name):
        
        if m_name == 0:
            stk_names = meas.configuration(m_name).stack_names()
            stk = meas.stack(stk_names[0])
            pix_data = stk.data()
            im = np.mean(pix_data, axis=3)[0]
            fig = pp.figure(figsize=(6,6))
            pp.axis('off')
            pp.imshow(im, cmap ='hot' )
            pp.text(10, 30, 'Overview',color='white', fontsize=15, fontweight='bold')
            pp.savefig('D:/current data/testa.tiff',bbox_inches='tight')
            pp.close(fig)
            document.add_picture('D:/current data/testa.tiff', width=Inches(5.5))
        else:
            stk_names = meas.configuration(m_name).stack_names()
            stk_g1 = meas.stack(stk_names[0])
            stk_g2 = meas.stack(stk_names[2])
            stk_r1 = meas.stack(stk_names[1])
            stk_r2 = meas.stack(stk_names[3])
            
            pix_data_green = stk_g1.data() + stk_g2.data()
            pix_data_red = stk_r1.data() + stk_r2.data()
            im_g = np.mean(pix_data_green, axis=1)[0]
            im_r = np.mean(pix_data_red, axis=1)[0]
            
                
            fig = pp.figure(figsize=(6,6))
            pp.subplot(221)
            pp.imshow(im_g, cmap ='hot' )
            pp.axis('off')
            pp.text(3, 10, '{}{}'.format(meas_names[m_name],': green'),color='white', fontsize=15, fontweight='bold')
            pp.subplot(222)
            pp.imshow(im_r, cmap ='hot' )
            pp.axis('off')
            pp.text(3, 10, '{}{}'.format(meas_names[m_name],': red'),color='white', fontsize=15, fontweight='bold')
            pp.savefig('D:/current data/testa.tiff',bbox_inches='tight')
            pp.close(fig)
            document.add_picture('D:/current data/testa.tiff', width=Inches(5.5))
            
            
    
    
    
    def test(dicta, levels):
            
            
            if type(dicta) == type(True) or type(dicta) == int or type(dicta) == list():
                print('level0')
            else:
                for i in dicta.keys():
                    style = document.styles['Normal']
                    paragraph_format = style.paragraph_format
                    
                    document.add_page_break()
                    p1_header = document.add_heading(str(i), level=1)
                    p1_header.paragraph_format.left_indent = Inches(0.0)
                    
                    if type(dicta[i]) == int or type(dicta[i]) == type(True) or type(dicta[i]) == str :#or type(dicta[i]) == list:
                        p1 = document.add_paragraph(str(dicta[i]))
                        p1.paragraph_format.left_indent = Inches(0.5)
                        
                    elif type(dicta[i]) == dict:
                        for ii in dicta[i].keys():
                            p2_header = document.add_heading(str(ii),level = 2)
                            p2_header.paragraph_format.left_indent = Inches(0.5)
                            if type(dicta[i][ii]) == int or type(dicta[i][ii]) == type(True) or type(dicta[i][ii]) == str or type(dicta[i][ii]) == list or type(dicta[i][ii]) == float:
                                p2 = document.add_paragraph(str(dicta[i][ii]))
                                p2.paragraph_format.left_indent = Inches(0.5)
                                
                            else:
                                for iii in dicta[i][ii].keys():
                                    p3_header = document.add_heading(str(iii),level = 3)
                                    p3_header.paragraph_format.left_indent = Inches(1)
                                    if type(dicta[i][ii][iii]) == int or type(dicta[i][ii][iii]) == type(True) or type(dicta[i][ii][iii]) == str or type(dicta[i][ii][iii]) == list or type(dicta[i][ii][iii]) == float:
                                           p3 = document.add_paragraph(str(dicta[i][ii][iii]))
                                           p3.paragraph_format.left_indent = Inches(1)
                                           
                                    else:
                                        for iiii in dicta[i][ii][iii].keys():
                                            p4_header = document.add_heading(str(iiii),level = 4)
                                            p4_header.paragraph_format.left_indent = Inches(2)
                                            p4 =document.add_paragraph(str(dicta[i][ii][iii][iiii]))
                                            p4.paragraph_format.left_indent = Inches(2)
                                            
    
    document = Document()
    
    im = specpy.Imspector()
    #meas= im.active_measurement()
    meas = im.measurement(im.measurement_names()[1])
    a = meas.parameter('')
    
    document.add_heading('{}{}'.format('Measurement-Report:                                                ',a['Measurement']['MeasTime']), 0)
    
    document.add_heading('Important Parameter', level=2)
    table = document.add_table(rows=6, cols=2)
    cell = table.cell(0,0)
    cell.text = 'Parameter'
    cell = table.cell(0,1)  
    cell.text = 'Value'
    
    meas_names = meas.configuration_names()
    
    document.add_page_break()
    for ii in range(len(meas_names)):
        GT(ii)    
    table.style = 'Light Grid Accent 1'
    test(a,0)                   
    document.save('{}{}'.format(path,'/Meas_protocol.docx'))
    
    
           
        
        
def MultiRun_meas(): 
    import time
    import math
    
    runs = int(MultiRUN_01.get()) ### define the number of Rois you want to scan
    aa  = 20  ### here place the ROISize
    delaytime = 1 # here we slow down the loop in [s] in order to make sure all Rois will be identified
    
                                                  
    im = specpy.Imspector()
    meas= im.active_measurement()
    
    #x_off = 1 # c.parameters('ExpControl/scan/range/offsets/coarse/x/g_off') # current stage position in x [m]
    y_off = meas.parameters('ExpControl/scan/range/offsets/coarse/y/g_off') # current stage position in x [m]
    y_add = y_off + 1.1*roi_size * np.linspace(0,runs-1, runs) ### initial position 
    
                                        
    z_position =  5e-08 #float(pixelsize)*1e-09         # in meter
    x_pixelsize = float(pixelsize)*1e-09        # in meter
    y_pixelsize = float(pixelsize)*1e-09        # in meter
    z_pixelsize = float(pixelsize)*1e-09        # in meter
    RZ =    float(Roisize)*1e-06            # in meter    
    Dwelltime= float(dwelltime)*1e-06         # in seconds 
    number_frames = float(frame_number) 
                                
    time_wait = math.ceil((RZ/x_pixelsize) * (RZ/y_pixelsize)* number_frames* Dwelltime) 
    print(time_wait)
    
    
    
    for i in y_add:
        print('Overview=',i)
        
        Overview(1,i)
        #time.sleep(5)
        Findpeak()
        time.sleep(1)
        Run_meas(1,i) ### 1 = TRUE for MUltirun
        
        print('sample=',y_add,'#peaks=',number_peaks_new, 'timewait=',(time_wait * number_peaks_new +1))#delaytime*number_peaks_new)
        time.sleep((time_wait * number_peaks_new))
        
        SAVING(save_path)
        im.close(im.measurement(im.measurement_names()[1]))
        
    
    
    
    
    
button_1 = Button(labtext_1, width = 10, text = 'Connect',    
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = Connect, anchor = 'w').grid(row = 0, column = 0)
button_2 = Button(labtext_1, width = 9, text = 'Overview',
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = partial(Overview,0,0), anchor = 'w').grid(row = 0, column = 1)
button_3 = Button(labtext_1, width = 9, text = 'FindPeak',   
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = Findpeak).grid(row = 0, column = 2)
button_4 = Button(labtext_1, width = 9, text = 'Run',   
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = partial(Run_meas,0,0)).grid(row = 0, column = 3)
button_5 = Button(labtext_1, width = 10, height =1, text = 'Set value',   
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = TEST).grid(row = 1, column = 0)
button_6 = Button(labtext_1, width = 9, height =1, text = 'Power',   
                  activebackground= 'green',font = ('Sans','9','bold'),
                  activeforeground= 'red', command = powerseries).grid(row = 1, column = 1)
button_7 = Button(labtext_1, width = 9, height =1, text = 'Pinhole',   
                  activebackground= 'green', font = ('Sans','9','bold'),
                  activeforeground= 'red', command = pinholeseries).grid(row = 1, column = 2)
button_8 = Button(labtext_1, width = 9, height =1, text = 'MultiRun',   
                  activebackground= 'green', font = ('Sans','9','bold'),
                  activeforeground= 'red', command = MultiRun_meas).grid(row = 1, column = 3)


scale_01_label= Label(frame_7, text='Thres:', height = 1,foreground= 'white', background =colour, font = ('Sans','9','bold'))
scale_01_label.grid(row = 0, column = 0, sticky = W+N)
scale_01 = Scale(frame_7, from_=0.5, to=5, showvalue=1, background =colour)
scale_01_value = scale_01.get()
scale_01.set(20)
scale_01.bind("<ButtonRelease-1>", SSS)
scale_01.grid(row = 1, column = 0, sticky = W+N)

scale_02_label= Label(frame_7, text='Rmin:', height = 1,foreground= 'white', background =colour, font = ('Sans','9','bold'))
scale_02_label.grid(row = 0, column = 1, sticky = W+N)
scale_02 = Scale(frame_7, from_=1, to=50, showvalue=1, background =colour)
scale_02_value = scale_02.get()
scale_02.set(50)
scale_02.bind("<ButtonRelease-1>", SSS)
scale_02.grid(row = 1, column = 1, sticky = W+N)

scale_03_label= Label(frame_7, text='Rmax:', height = 1,foreground= 'white', background =colour, font = ('Sans','9','bold') )
scale_03_label.grid(row = 0, column = 2, sticky = W+N)
scale_03 = Scale(frame_7, from_=5, to=1000, showvalue=1, background =colour)
scale_03_value = scale_03.get()
scale_03.set(1000)
scale_03.bind("<ButtonRelease-1>", SSS)
scale_03.grid(row = 1, column = 2, sticky = W+N)


    
root.mainloop()
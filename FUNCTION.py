# -*- coding: utf-8 -*-
"""
Created on Tue May 19 15:41:20 2020

@author: Abberior_admin

edits by NV
this code has numerous serious issues
it works, but it contains several features that are considered bad practice

unfixed:   
several pieces of code are repeated multiple times - they should go in functions
there are a lot of pieces whose function seem unnesecary.
there are many typecasts without changing the variable name
as a solution one might:
going through the whole of the code to much-out unneeded parts.

done:
it had wildcard import - now fixed
namespace names are overwritten with variable names, e.g.
  import impspector as im; im = Image.image() - fixed
a consequence is that in each function namespaces are reloaded
global variables are used at random places - fixed
button handles and variables are passed in long lists of function returns
  making it easy to mistake the order. - fixed
As a solution one might:
build the app into a class, solving the need for global variables and
  passing so many variables
"""
from PIL import Image,ImageTk
from tkinter import ttk
import specpy
import numpy as np
import matplotlib.pyplot as pp
import os
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg#, NavigationToolbar2TkAgg
import matplotlib.cm as cm
import tkinter as tk
import time
from datetime import datetime
import threading  
from pprint import pprint # for debugging
import plotTrace # homebuild to plot recorded t-traces
import GUISpotFinding as spotFinding


    
def Connect(self):
    T = self.T
    try:
        im=specpy.Imspector()
        T.insert(tk.END, 'connection successful\n')
    except (RuntimeError):
        
        T.insert(tk.END, 'connection failed\n')    
    return

def makeOverview(self,Multi, Pos):
    """creates an overview image for finding suitable spots,
    parts of this function can be replaced by calling
    the method applyGUISettings, but this is not implemented yet"""

    
    #get values from GUI entries
    roi_size = float(self.ROIsize_overview_value.get())*1e-06          # in meter
    Dwelltime= np.around(float(self.dwell_overview_value.get()))*1e-06         # in seconds  
    number_frames = int(self.frames_overview_value.get())      #Type number of frames t in xyt mode
    #z_position = 0        # in meter
    x_pixelsize = float(self.pxsize_overview_value.get())*1e-09       # in meter
    y_pixelsize = float(self.pxsize_overview_value.get())*1e-09       # in meter
    z_pixelsize = float(self.pxsize_overview_value.get())*1e-09       # in meter
    px_num = roi_size/x_pixelsize
    
     #here the laser values are read in, multiple simultaneous can be indicated split by kommas
    laser_overview = [int(s) for s in self.laser_overview_entry.get().split(',')]
    laser_overview_VALUE = [int(s) for s in self.laser_overview_value.get().split(',')]
    
    #translate string from GUI to position in array
    lasersOn = []
    for i in range(len(laser_overview)):
        if laser_overview[i]  == 485:laservalue = 0
        elif laser_overview[i]  == 518:laservalue = 1
        elif laser_overview[i]  == 595:laservalue = 2
        elif laser_overview[i]  == 561:laservalue = 3
        elif laser_overview[i]  == 640:laservalue = 4
        elif laser_overview[i]  == 775:laservalue = 5
        lasersOn.append(laservalue)
    LASER_ACT =[False]*8
    for laser in lasersOn:
        LASER_ACT[laser] = True
    
    ############
    
    path = 'D:/current data/'
    im = specpy.Imspector()#re-get connection
    msr=im.create_measurement()
    config = msr.active_configuration()
    # overview is called multiple times to identify spot locations for the next spot.
    if Multi == 1: 
        print('MULTIRUN')
        config.set_parameters('ExpControl/scan/range/offsets/coarse/y/g_off', Pos) # current stage position in x [m]
    else:
        print('non MULTIRUN')
    #must be set first such that settings are available? no.
    xyt_mode = 784# for xyt mode  784 #for xyz_mode 528
    config.set_parameters('ExpControl/scan/range/mode',xyt_mode) 

    #a bunch of these steps are default and can be put into a helper function
    applyScannerSettings(self, config)
    applyDetectorSettings(self, config, 
                        enableStream = False, 
                        stream = 'HydraHarp')
    #roi size, pixel size and laser settings are first set in above functions
    #then overwritten herer
    config.set_parameters('ExpControl/scan/range/x/psz',x_pixelsize)
    config.set_parameters('ExpControl/scan/range/y/psz',y_pixelsize)
    config.set_parameters('ExpControl/scan/range/z/psz',z_pixelsize)
    config.set_parameters('ExpControl/scan/range/x/len',roi_size)
    config.set_parameters('ExpControl/scan/range/y/len',roi_size)
    config.set_parameters('ExpControl/scan/dwelltime', Dwelltime)
    config.set_parameters('ExpControl/scan/range/t/res',number_frames)
    #use only one linestep
    config.set_parameters('ExpControl/gating/linesteps/steps_active',
                          [True, False, False, False, False, False, False, False])
    config.set_parameters('ExpControl/gating/linesteps/laser_on',[LASER_ACT]*8)
    #apply laser settings
    for laser, power in zip(lasersOn, laser_overview_VALUE):
        config.set_parameters('ExpControl/lasers/power_calibrated/%i/value/calibrated' % laser, power)
    
    #here the measurement is actually run
    im.run(msr)

    #here the data is read out of the image
    #get data and sum over the axes t and z
    xy_data = [np.sum(config.stack(i).data(), axis = (0, 1)) for i in range(4)]
    self.xy_data = xy_data
    datashape = config.stack(0).data().shape
    print('xy dimensions of overview image is (%i, %i)'% xy_data[0].shape)
    
    #the data array should have format [time, z, y, x]
    assert datashape == (1,number_frames,np.round(px_num,0),np.round(px_num,0)), \
        'The image shape does not match the expected shape'
        
    #from all channels, the relevant ones are added to make the overview image
    #peak selection also occurs on the overview image
    peakchannels = self.peakchannels.get()
    firstchannel = int(peakchannels[0])
    overview = self.xy_data[firstchannel]
    #add more channels if given
    for chan in peakchannels[1:]:
        overview += self.xy_data[int(chan)]

    #display overview image, see 
    #https://stackoverflow.com/questions/10965417/how-to-convert-a-numpy-array-to-pil-image-applying-matplotlib-colormap
    photo = cm.hot(overview.astype(np.float)/max(overview.flatten())) # cm needs floats
    photo = np.uint8(photo * 255)
    photo = Image.fromarray(photo)
    photo = photo.resize((400, 400), Image.ANTIALIAS)
    photoTk = ImageTk.PhotoImage(photo)
    photo.save('{}{}'.format(path,'Overview_image.tiff'), format = 'TIFF')#'D:/current data/','Overview_image.tiff'
    label = tk.Label(self.frame_topleft, image=photoTk)
    #this seems unnesecary, but for some reason it is needed
    label.image = photoTk
    label.grid(row=0)
    self.T.delete('1.0', tk.END) 
    self.T.insert(tk.END, "Overview created\n") 
    return
    
# def Findpeak(self):
#     """currently a dummy that returns 3 random xy pairs"""
#     import random
#     self.roi_xs = np.array([random.random() *10 for i in range(3)])
#     self.roi_ys = np.array([random.random() *10 for i in range(3)])
#     self.number_peaks = len(self.roi_xs)

    
def _Run_meas(*args):  
    """this function splits off Run meas into a separate thread, such
    that the GUI remains responsive and the run may be aborted"""
    thread = threading.Thread(target=Run_meas, args = args)  
    thread.start()  

def Run_meas(self):
    """
    reads a list of positions attached to the class by self.roix and self.roiy
    for each position an image if recorded according to the settings in the GUI
    the data is saved as a ptu file, an overview image of the region is also saved
    
    Large portions of this function can be replaced by the method 
    applyGUISettings. TODO

    Returns
    -------
    None.

    """
    Pos = self.y_coarse_offset
    pixelsize = self.pxsize.get()
    #values are by default read in as tring, have to convert
    pixelsize_global = float(self.pxsize_overview_value.get()) * 1e-9
    
    #consider moving below part to separate function
    #z_position =  5e-08 #float(pixelsize)*1e-09         # in meter
    x_pixelsize = float(pixelsize)*1e-09        # in meter
    y_pixelsize = float(pixelsize)*1e-09        # in meter
    z_pixelsize = float(pixelsize)*1e-09        # in meter
    roi_size =    float(self.ROIsize.get())*1e-06            # in meter 
    Dwelltime= float(self.dwelltime.get())*1e-06         # in seconds  
           
    LP485 = self.L485_value.get()
    LP518 = self.L518_value.get()
    LP561 = self.L561_value.get()
    LP640 = self.L640_value.get()
    #LP595 = self.L595_value.get()
    LP775 = self.L775_value.get()
                
    Streaming_HydraHarp= True # Type True or False for Streaming via HydraHarp 
    modelinesteps= True      #Type True for linesteps
    number_linesteps = 2     #Type the number of linesteps
    xyt_mode = 784           #Type 785 for xyt mode
    
    number_frames = float(self.NoFrames.get())     #Type number of frames t in xyt mode
    #time_wait = math.ceil((roi_size/x_pixelsize) * (roi_size/y_pixelsize)* number_frames* Dwelltime) + 1
    #measurement time consists of line scan rate + flyback time + buffer
    time_wait = 1
                    
    Activate485 = bool(self.L485_1)                    
    Activate518 = bool(self.L518_1)   
    Activate561 = bool(self.L561_1)   
    Activate640 = bool(self.L640_1)   
    Activate595 = bool(self.L595_1)   
    Activate775 = bool(self.L775_1)   
    
    Activate485_02 = bool(self.L485_2)                    
    Activate518_02 = bool(self.L518_2)   
    Activate561_02 = bool(self.L561_2)   
    Activate640_02 = bool(self.L561_2)   
    Activate595_02 = bool(self.L595_2)   
    Activate775_02 = bool(self.L775_2) 
    #Activate_Autofocus= bool(act_Autofocus)
        
        
            
        ############################
        #Hydraharp-setting
            
    if Streaming_HydraHarp == True:
        stream = 'HydraHarp'
    else:
        stream = 'fpga'
                       
    Ch1_mode= 0                # type 0 for counter and 1 for flim
    Ch1_stream= stream         # type hydraharp or fpga for internal
    Ch2_mode= 0          
    Ch2_stream= stream
    Ch3_mode= 0         
    Ch3_stream= stream
    Ch4_mode= 0         
    Ch4_stream= stream
            
        ############################
        #detector activation
        
    detector1= True
    detector2= True
    detector3= True
    detector4= True
            
       ####################################################################################
            
        
    linesteps = [False, False, False, False, False, False, False, False]
    for i in range(number_linesteps):
        linesteps[i] = True
            
    
    im = specpy.Imspector() 
    d = im.active_measurement()
    M_obj= im.measurement(d.name())
             
#    if Activate_Autofocus == True: #Activate_Autofocus
#        M_obj.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', False)
#        time.sleep(2) 
#        M_obj.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', True)
        
    try:
        x_roi_new = self.roi_xs
        y_roi_new = self.roi_ys 
    except RecursionError:
        self.T.insert(tk.END, 'no peaks positions in memory')
        raise RecursionError
    #location to save the files collected in next loop
    dateTimeObj = datetime.now()
    timestamp = dateTimeObj.strftime("%Y-%b-%d-%H-%M-%S")
    save_path = os.path.join(self.dataout, timestamp + 'Overview_%.2f_numberSPOTS_%i' % (Pos, len(self.roi_xs)))
    try:
        os.mkdir(save_path)
    except FileExistsError:
        self.T.insert(tk.END, 'saving directory already exists, skipping')
    for i in range(x_roi_new.size):
        x_position = x_roi_new[i]
        y_position = y_roi_new[i]
        d = im.active_measurement()
        M_obj.clone(d.active_configuration())
        # this must not be the i-th but the latest configuration
        M_obj.activate(M_obj.configuration(i))          
        c = M_obj.configuration(i)
        applyLaserSettings(self, c)
        
        c.set_parameters('ExpControl/scan/range/x/psz',x_pixelsize)
        c.set_parameters('ExpControl/scan/range/y/psz',y_pixelsize)
        c.set_parameters('ExpControl/scan/range/z/psz',z_pixelsize)
        c.set_parameters('ExpControl/scan/range/x/off', x_position*pixelsize_global )#+ ROI_offset)
        c.set_parameters('ExpControl/scan/range/y/off', y_position*pixelsize_global )#+ ROI_offset)
        c.set_parameters('ExpControl/scan/range/z/off', 1e-15)#z_position*z_pixelsize)
        c.set_parameters('ExpControl/scan/range/x/len',roi_size)
        c.set_parameters('ExpControl/scan/range/y/len',roi_size)
        c.set_parameters('ExpControl/scan/range/z/len',roi_size)
        c.set_parameters('ExpControl/scan/dwelltime', Dwelltime)
        c.set_parameters('HydraHarp/data/streaming/enable', Streaming_HydraHarp)
        c.set_parameters('HydraHarp/is_active', Streaming_HydraHarp)
        c.set_parameters('ExpControl/gating/tcspc/channels/0/mode', Ch1_mode)
        c.set_parameters('ExpControl/gating/tcspc/channels/0/stream', Ch1_stream)
        c.set_parameters('ExpControl/gating/tcspc/channels/1/mode', Ch2_mode)
        c.set_parameters('ExpControl/gating/tcspc/channels/1/stream', Ch2_stream)
        c.set_parameters('ExpControl/gating/tcspc/channels/2/mode', Ch3_mode)
        c.set_parameters('ExpControl/gating/tcspc/channels/2/stream', Ch3_stream)
        c.set_parameters('ExpControl/gating/tcspc/channels/3/mode', Ch4_mode)
        c.set_parameters('ExpControl/gating/tcspc/channels/3/stream', Ch4_stream)
        c.set_parameters('ExpControl/gating/linesteps/on', modelinesteps)
        c.set_parameters('ExpControl/gating/linesteps/steps_active', linesteps)
        c.set_parameters('ExpControl/gating/linesteps/step_values',[1,1,0,0,0,0,0,0])
        c.set_parameters('ExpControl/scan/range/mode',xyt_mode)
        c.set_parameters('ExpControl/lasers/power_calibrated/0/value/calibrated', float(LP485))
        c.set_parameters('ExpControl/lasers/power_calibrated/2/value/calibrated', float(LP518))
        c.set_parameters('ExpControl/lasers/power_calibrated/3/value/calibrated', float(LP561))
        c.set_parameters('ExpControl/lasers/power_calibrated/4/value/calibrated', float(LP640))
        c.set_parameters('ExpControl/lasers/power_calibrated/5/value/calibrated', float(LP775))
        c.set_parameters('ExpControl/gating/pulses/pulse_chan/delay',[0.0, 0.0, 0.0, 0.0])
        c.set_parameters('ExpControl/gating/linesteps/laser_enabled',[True, True, True, True, True, True, False, False])

        c.set_parameters('ExpControl/scan/range/t/res',number_frames)
        c.set_parameters('ExpControl/gating/linesteps/chans_enabled',[detector1,detector2,detector3,detector4])
        c.set_parameters('ExpControl/scan/detsel/detsel',['APD1', 'APD2', 'APD3', 'APD4'])
        c.set_parameters('ExpControl/gating/linesteps/chans_on', [[True, True, True, True],
         [True, True, True, True],
         [False, False, False, False],
         [False, False, False, False],
         [False, False, False, False],
         [False, False, False, False],
         [False, False, False, False],
         [False, False, False, False]])
           

        im.run(M_obj)
        time.sleep(time_wait) 
        
        if self.abort_run:
            self.T.insert(tk.END, 'aborting (multi-) measurement run\n')
            self.T.insert(tk.END, 'be sure to reset abort before your next run\n')
            break
                
    #save msr file, move files to subfolder
    msrout = os.path.join(save_path, 'Overview%.2f.msr' % Pos)
    im.measurement(im.measurement_names()[1]).save_as(msrout)
     
    files = os.listdir('D:/current data/')
    files_ptu = [i for i in files if i.endswith('.ptu')]
    files_dat = [i for i in files if i.endswith('.dat')]
    
    n_files_ptu = len(files_ptu)
    n_files_dat = len(files_dat)
        
    for ii in range(n_files_ptu):
        os.rename('{}{}'.format('D:/current data/',files_ptu[ii]), \
                  '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, '_spot_', ii, '.ptu'))
    
    for ii in range(n_files_dat):
        os.rename('{}{}'.format('D:/current data/',files_dat[ii]), \
                  '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, \
                                          '_spot_', files_dat[ii],'.dat'))
    return save_path

def _timeRun(*args):  
    """this function splits off Run meas into a separate thread, such
    that the GUI remains responsive and the run may be aborted"""
    thread = threading.Thread(target=timeRun, args = args)  
    thread.start()  
def timeRun(self):
    Pos = self.y_coarse_offset
    #values are by default read in as tring, have to convert
    pixelsize = float(self.pxsize_overview_value.get()) * 1e-9
    # #Activate_Autofocus= bool(act_Autofocus)   

    # I don't understand why this is needed, or how it works,
    #but removing it causes strange behaviour         
    im = specpy.Imspector() 
    msr = im.active_measurement()
    #msr = im.measurement(d.name()) 
#    if Activate_Autofocus == True: #Activate_Autofocus
#        config.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', False)
#        time.sleep(2) 
#        config.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', True)

    try:
        x_roi_new = self.roi_xs
        y_roi_new = self.roi_ys 
    except RecursionError:
        self.T.insert(tk.END, 'no peaks positions in memory')
        raise RecursionError
        
    #save the overview image and to be imaged spots
    spotFinding.plotpeaks(self.smoothimage, self.goodpeaks,\
                          savedir = self.dataout, isshow = True)
    acquisition_time = 3
    for i in range(x_roi_new.size):
        print("analysing spot %i out of %i" % (i, len(x_roi_new)))
        x_position = x_roi_new[i]
        y_position = y_roi_new[i]
        #d = im.active_measurement()
        msr.clone(msr.active_configuration())



        #seems that the cloned measurement should be automatically activated, need to test
        config = msr.active_configuration()
# =============================================================================
#        # the last config name corresponds to the newest configuration
#         config_name = msr.configuration_names()[-1]
#         config = msr.configuration(config_name)
#         # msr.activate(config)          
# =============================================================================
        applyLaserSettings(self, config)
        applyDetectorSettings(self, config, stream = 'HydraHarp',
                              enableStream = True)
        # for xyt mode  784    #for xyz_mode 528    #for t 1363
        config.set_parameters('ExpControl/scan/range/mode',1363)
        config.set_parameters('ExpControl/scan/range/t/len', acquisition_time)
        config.set_parameters('ExpControl/scan/range/x/off', x_position*pixelsize )
        config.set_parameters('ExpControl/scan/range/y/off', y_position*pixelsize )
        # config.set_parameters('ExpControl/gating/tcspc/channels/0/mode', 0) 
        # config.set_parameters('ExpControl/gating/tcspc/channels/0/stream', stream)
        # config.set_parameters('ExpControl/gating/tcspc/channels/1/mode', 0)
        # config.set_parameters('ExpControl/gating/tcspc/channels/1/stream', stream)
        # config.set_parameters('ExpControl/gating/tcspc/channels/2/mode', 0)
        # config.set_parameters('ExpControl/gating/tcspc/channels/2/stream', stream)
        # config.set_parameters('ExpControl/gating/tcspc/channels/3/mode', 0)
        # config.set_parameters('ExpControl/gating/tcspc/channels/3/stream', stream)
        # config.set_parameters('HydraHarp/data/streaming/enable', True)
        # config.set_parameters('HydraHarp/is_active', True) #not sure what this does
        im.run(msr)
        print("finished position x %i and y %i" % (x_position, y_position))

        #if by accident the integration time is really short, this
        #should prevent the software from crashing
        if acquisition_time < 1:
            time.sleep(2) 
        
        
        #plot t trace of last ptu file
        lastfile = plotTrace.getLastModified(self.dataout)
        channels = plotTrace.getTraces(lastfile, [0,2])
        binneddata = plotTrace.plotTrace(channels, (0,acquisition_time), lastfile,\
                                         step = 5e-3, outname = 'inferred')

        if self.abort_run: #not Implemented, need to split function off in seperate
        #thread like in _run_meas
            self.T.insert(tk.END, 'aborting (multi-) measurement run\n')
            self.T.insert(tk.END, 'be sure to reset abort before your next run\n')
            break
        
    #location to save the files collected in next loop
    dateTimeObj = datetime.now()
    timestamp = dateTimeObj.strftime("%Y-%b-%d_%H-%M-%S")
    save_path = os.path.join(self.dataout, timestamp + 'Overview_%.2f_numberSPOTS_%i' % (Pos, len(self.roi_xs)))
    try:
        os.mkdir(save_path)
    except FileExistsError:
        self.T.insert(tk.END, 'saving directory already exists, skipping')
        
    #save msr file, move files to subfolder
    msrout = os.path.join(save_path, 'Overview%.2f.msr' % Pos)
    msr.save_as(msrout)
     
    #shift files to save dir
    files = os.listdir(self.dataout)
    extensions = ['.png', '.dat', '.ptu', '.txt', 'tiff']
    files_selected = [i for i in files if i[-4:] in extensions]    
    for file in files_selected:
        source = os.path.join(self.dataout, file)
        dest = os.path.join(save_path, file)
        os.rename(source, dest)
    #I don't like the renaming of spots in this way. timestamp is better
# =============================================================================
#     for ii in range(n_files_ptu):
#         os.rename('{}{}'.format('D:/current data/',files_ptu[ii]), \
#                   '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, '_spot_', ii, '.ptu'))
#     
#     for ii in range(n_files_dat):
#         os.rename('{}{}'.format('D:/current data/',files_dat[ii]), \
#                   '{}{}{}{}{}{}{}'.format(save_path,'/','Overview_Pos_y', Pos, \
#                                           '_spot_', files_dat[ii],'.dat'))
# =============================================================================
    return save_path

def applyScannerSettings(self, config):
    """for most of the measurement a lot of settings are default, they are set
    for the meas object that is passed to this function
    the time measurement was giving some trouble, now it is split somewhat, have to clean up"""
    
    
    Dwelltime= float(self.dwelltime.get())*1e-06         # in seconds
    pixelsize = float(self.pxsize.get())*1e-09          # in meter
    roi_size =    float(self.ROIsize.get())*1e-06            # in meter
    number_frames = float(self.NoFrames.get())     #Type number of frames t in xyt mode
    #this seems to enable the autofocus, is it really so?
    config.set_parameters('OlympusIX/scanrange/z/z-stabilizer/enabled', True)
    config.set_parameters('ExpControl/scan/range/z/off', 1e-15)#z_position*z_pixelsize)
    
    #setting the pinhole here silently? need to communicate this clearly
    config.set_parameters('Pinhole/pinhole_size', 10e-5)
    
    #set the scanner size
    config.set_parameters('ExpControl/scan/range/x/psz',pixelsize)
    config.set_parameters('ExpControl/scan/range/y/psz',pixelsize)
    #xy offset is set in the function
    config.set_parameters('ExpControl/scan/range/x/len',roi_size)
    config.set_parameters('ExpControl/scan/range/y/len',roi_size)
    config.set_parameters('ExpControl/scan/range/z/len',0)
    config.set_parameters('ExpControl/scan/dwelltime', Dwelltime)
    config.set_parameters('ExpControl/scan/range/t/res',number_frames)
def applyDetectorSettings(self, config, 
                        enableStream = False, 
                        stream = 'HydraHarp'):
    assert stream == 'HydraHarp' or stream == 'fpga', 'bad stream value'
    config.set_parameters('HydraHarp/is_active', True) #not sure what this does
    config.set_parameters('HydraHarp/data/streaming/enable', enableStream)
    # mode codes: type 0 for counter and 1 for flim
    config.set_parameters('ExpControl/gating/tcspc/channels/0/mode', 0) 
    config.set_parameters('ExpControl/gating/tcspc/channels/0/stream', stream)
    config.set_parameters('ExpControl/gating/tcspc/channels/1/mode', 0)
    config.set_parameters('ExpControl/gating/tcspc/channels/1/stream', stream)
    config.set_parameters('ExpControl/gating/tcspc/channels/2/mode', 0)
    config.set_parameters('ExpControl/gating/tcspc/channels/2/stream', stream)
    config.set_parameters('ExpControl/gating/tcspc/channels/3/mode', 0)
    config.set_parameters('ExpControl/gating/tcspc/channels/3/stream', stream)
    config.set_parameters('ExpControl/scan/detsel/detsel',['APD1', 'APD2', 'APD3', 'APD4']) # is this needed?
    config.set_parameters('ExpControl/gating/linesteps/chans_enabled',[True, True, True, True]) # what does this do?
    config.set_parameters('ExpControl/gating/pulses/pulse_chan/delay',[0.0, 0.0, 0.0, 0.0]) # not sure if needed
    #where can I set how many linesteps are on?
    config.set_parameters('ExpControl/gating/linesteps/laser_enabled',[True, True, True, True, True, True, False, False])
    config.set_parameters('ExpControl/gating/linesteps/chans_on', \
                        [[True]*4, [True]*4, [False]*4, [False]*4, \
                         [False]*4, [False]*4, [False]*4, [False]*4])
       
def applyLaserSettings(self, config):
    #read-out GUI values in SI units

    LP485 = float(self.L485_value.get()) # in %
    LP518 = float(self.L518_value.get())
    LP561 = float(self.L561_value.get())
    LP640 = float(self.L640_value.get())
    LP595 = float(self.L595_value.get()) #which id in power_calibrated does this have? Somewhere it was written that it should be 2.
    LP775 = float(self.L775_value.get())
    Activate485 = bool(self.L485_1)                    
    Activate518 = bool(self.L518_1)   
    Activate561 = bool(self.L561_1)   
    Activate640 = bool(self.L640_1)   
    Activate595 = bool(self.L595_1)   
    Activate775 = bool(self.L775_1)   
    Activate485_02 = bool(self.L485_2)                    
    Activate518_02 = bool(self.L518_2)   
    Activate561_02 = bool(self.L561_2)   
    Activate640_02 = bool(self.L640_2)   
    Activate595_02 = bool(self.L595_2)   
    Activate775_02 = bool(self.L775_2) 
    config.set_parameters('ExpControl/lasers/power_calibrated/0/value/calibrated', float(LP485))
    config.set_parameters('ExpControl/lasers/power_calibrated/2/value/calibrated', float(LP518))
    config.set_parameters('ExpControl/lasers/power_calibrated/3/value/calibrated', float(LP561))
    config.set_parameters('ExpControl/lasers/power_calibrated/4/value/calibrated', float(LP640))
    config.set_parameters('ExpControl/lasers/power_calibrated/5/value/calibrated', float(LP775))
    config.set_parameters('ExpControl/gating/linesteps/laser_on',[[Activate485, Activate518, Activate595, Activate561, Activate640, Activate775, False, False],
         [Activate485_02, Activate518_02, Activate595_02, Activate561_02, Activate640_02, Activate775_02, False, False],
         [False]*8,[False]*8,[False]*8,[False]*8,[False]*8,[False]*8]) #the last six linesteps are not used
    
def SAVING(path,a):
    from docx import Document
    #from docx.shared import Pt
    #from docx.shared import Length
    #from docx.shared import RGBColor
    #from docx.enum.text import WD_LINE_SPACING
    #from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    #import matplotlib as pp
    
    if a==0:
        #### functions #####
        def GT(m_name):
            import matplotlib as pp
            
            if m_name == 0:
                stk_names = meas.configuration(m_name).stack_names()
                stk = meas.stack(stk_names[0])
                pix_data = stk.data()
                im = np.mean(pix_data, axis=3)[0]
                fig = pp.figure(figsize=(6,6))
                pp.axis('off')
                pp.imshow(im, cmap ='hot' )
                pp.text(10, 30, 'Overview',color='white', fontsize=15, fontweight='bold')
                pp.savefig('D:/current data/testa.tiff',bbox_inches='tight')
                pp.close(fig)
                document.add_picture('D:/current data/testa.tiff', width=Inches(5.5))
            else:
                stk_names = meas.configuration(m_name).stack_names()
                stk_g1 = meas.stack(stk_names[0])
                stk_g2 = meas.stack(stk_names[2])
                stk_r1 = meas.stack(stk_names[1])
                stk_r2 = meas.stack(stk_names[3])
                
                pix_data_green = stk_g1.data() + stk_g2.data()
                pix_data_red = stk_r1.data() + stk_r2.data()
                im_g = np.mean(pix_data_green, axis=1)[0]
                im_r = np.mean(pix_data_red, axis=1)[0]
                
                    
                fig = pp.figure(figsize=(6,6))
                pp.subplot(221)
                pp.imshow(im_g, cmap ='hot' )
                pp.axis('off')
                pp.text(3, 10, '{}{}'.format(meas_names[m_name],': green'),color='white', fontsize=15, fontweight='bold')
                pp.subplot(222)
                pp.imshow(im_r, cmap ='hot' )
                pp.axis('off')
                pp.text(3, 10, '{}{}'.format(meas_names[m_name],': red'),color='white', fontsize=15, fontweight='bold')
                pp.savefig('D:/current data/testa.tiff',bbox_inches='tight')
                pp.close(fig)
                document.add_picture('D:/current data/testa.tiff', width=Inches(5.5))
                
                
        
        
        
        def test(dicta, levels):
                
                
                if type(dicta) == type(True) or type(dicta) == int or type(dicta) == list():
                    print('level0')
                else:
                    for i in dicta.keys():
                        style = document.styles['Normal']
                        paragraph_format = style.paragraph_format
                        
                        document.add_page_break()
                        p1_header = document.add_heading(str(i), level=1)
                        p1_header.paragraph_format.left_indent = Inches(0.0)
                        
                        if type(dicta[i]) == int or type(dicta[i]) == type(True) or type(dicta[i]) == str :#or type(dicta[i]) == list:
                            p1 = document.add_paragraph(str(dicta[i]))
                            p1.paragraph_format.left_indent = Inches(0.5)
                            
                        elif type(dicta[i]) == dict:
                            for ii in dicta[i].keys():
                                p2_header = document.add_heading(str(ii),level = 2)
                                p2_header.paragraph_format.left_indent = Inches(0.5)
                                if type(dicta[i][ii]) == int or type(dicta[i][ii]) == type(True) or type(dicta[i][ii]) == str or type(dicta[i][ii]) == list or type(dicta[i][ii]) == float:
                                    p2 = document.add_paragraph(str(dicta[i][ii]))
                                    p2.paragraph_format.left_indent = Inches(0.5)
                                    
                                else:
                                    for iii in dicta[i][ii].keys():
                                        p3_header = document.add_heading(str(iii),level = 3)
                                        p3_header.paragraph_format.left_indent = Inches(1)
                                        if type(dicta[i][ii][iii]) == int or type(dicta[i][ii][iii]) == type(True) or type(dicta[i][ii][iii]) == str or type(dicta[i][ii][iii]) == list or type(dicta[i][ii][iii]) == float:
                                               p3 = document.add_paragraph(str(dicta[i][ii][iii]))
                                               p3.paragraph_format.left_indent = Inches(1)
                                               
                                        else:
                                            for iiii in dicta[i][ii][iii].keys():
                                                p4_header = document.add_heading(str(iiii),level = 4)
                                                p4_header.paragraph_format.left_indent = Inches(2)
                                                p4 =document.add_paragraph(str(dicta[i][ii][iii][iiii]))
                                                p4.paragraph_format.left_indent = Inches(2)
                                                
        
        document = Document()
        
        im = specpy.Imspector()
        meas = im.measurement(im.measurement_names()[1])
        a = meas.parameter('')    
        document.add_heading('{}{}'.format('Measurement-Report:                                                ',a['Measurement']['MeasTime']), 0) 
        document.add_heading('Important Parameter', level=2)
        table = document.add_table(rows=6, cols=2)
        cell = table.cell(0,0)
        cell.text = 'Parameter'
        cell = table.cell(0,1)  
        cell.text = 'Value'
        
        meas_names = meas.configuration_names()
        
        document.add_page_break()
        for ii in range(len(meas_names)):
            GT(ii)    
        table.style = 'Light Grid Accent 1'
        test(a,0)                   
        document.save('{}{}'.format(path,'/Meas_protocol.docx'))
    
    else:
        print('saving works')
        document = Document()
        document.add_heading('this is just a test document')
        document.save('{}{}'.format(path,'/Meas_protocol.docx'))
       
def layout(self):
    root = self.parent
    root.title("Imspector Control Interface")
    root1 = tk.Frame(root, width = 350, height = 350, bg = 'grey')
    root1.grid()
    icon_dir  = os.path.dirname(os.path.realpath(__file__))
    image_ID = Image.open(os.path.join(icon_dir,r'Alphatubulin.tif'))
    resized = image_ID.resize((400, 400), Image.ANTIALIAS)
    #when debugging, multiple tk instances can exist, then the master must be
    #specified for it to link the image with the right tcl instance.
    photo = ImageTk.PhotoImage(resized, master = root)
    colour = 'grey'
    txtcolour = 'white'
    space = 10
    framespacer= tk.Frame(root1, width = space, height = 2, bg = colour)
    
    
    frame_s1 = framespacer
    frame_s1.grid(row=0, column=0, sticky = 'n')
    frame_s2 = framespacer
    frame_s2.grid(row=0, column=3, sticky = 'n')
    frame_s3 = framespacer
    frame_s3.grid(row=4, column=0, sticky = 'n')
    frame_s4 = framespacer
    frame_s4.grid(row=4, column=3, sticky = 'n')
    frame_s5 = framespacer
    frame_s5.grid(row=2, column=1, sticky = 'n')
    frame_s6 = framespacer
    frame_s6.grid(row=2, column=2, sticky = 'n')
    frame_top = tk.Frame(root1, width = 400, height = 400, bg = colour)
    frame_top.grid(row=1, column=1)
    label = tk.Label(frame_top,image=photo)
    label.image = photo # this seems double, but it is actually needed, why?
    label.grid(row=0)
    
    frame_top2 = tk.Frame(root1, width = 400, height = 300, bg = colour)
    frame_top2.grid(row=3, column=1)
    frame_top2.grid_propagate(0)
    
    frame_top3 = tk.Frame(root1, width = 400, height = 400, bg = colour)
    frame_top3.grid(row= 1, column=2, sticky = 'n')
    
    frame_top4 = tk.Frame(root1, width = 400, height = 300, bg = colour)
    frame_top4.grid(row=3, column=2, sticky = 'n')

    f = pp.figure(figsize=(2,2), dpi=150, edgecolor='k')
    canvas = FigureCanvasTkAgg(f, master = frame_top3)
    canvas.get_tk_widget().grid(row=0, column=0)
    
    #this seems to have no function, but it causes trouble
    labtext_1 = tk.Frame(frame_top4,width = 300, height = 200,bg = colour)
    labtext_1.grid()
    
    
    T = tk.Text(frame_top4, height=10, width=37)
    T.grid()

    style = ttk.Style()
    settings = {"TNotebook.Tab": 
                   {"configure": {"padding": [1, 2], "background": colour , "font" : ('Sans','9','bold')}, 
                    "map": {"background": [("selected", "red"),  ("active", "#fc9292")], 
                            "foreground": [("selected", "pink"), ("active", "white")], 
                            "expand": [("selected", [1, 1, 1, 1])]} },
                "TFrame": 
                   {"configure": {"padding": [0, 2], "background": colour } },
                "TNotebook": 
                   {"configure": {"padding": [1, 2], "background": colour } }}
                
                       
                       
    style.theme_create('JHB', parent="alt", settings=settings)
    #style.theme_create('JHB2', parent="alt", settings=settings)
    style.theme_use('JHB')

    nb = ttk.Notebook(frame_top2, width = 7000)
    nb.pressed_index = None
    nb.grid(row=0, column=0, sticky='NESW')
    
    page1 = ttk.Frame(nb)
    nb.add(page1, text='Overview')
    
    page2 = ttk.Frame(nb)
    nb.add(page2, text='ROI_select')
    
    page3 = ttk.Frame(nb)   
    nb.add(page3, text='Powerseries')
    
    page4 = ttk.Frame(nb)
    nb.add(page4, text='Pinholeseries')
    
    page5 = ttk.Frame(nb)
    nb.add(page5, text='pointMeasurement')
    
    
    
    
    frame_spacer_01 = tk.Frame(page2, width = 50, height = 1, background =colour)
    frame_spacer_01.grid(row=0, column=0, sticky = 'w')
    
    frame_5 = tk.Frame(frame_spacer_01, width = 150, height = 110, background = colour)
    frame_5.grid(row=0, column=0, sticky = 'wn')
    frame_5.grid_propagate (False)
    
    frame_6 = tk.Frame(frame_spacer_01, width = 150, height = 180, background = colour)
    frame_6.grid(row=1, column=0, sticky = 'wn')
    frame_6.grid_propagate (False)
    
    frame_7 = tk.Frame(frame_spacer_01, width = 150, height = 180, background = colour)
    frame_7.grid(row=1, column=1, sticky = 'wn')
    frame_7.grid_propagate (False)
    
    frame_8 = tk.Frame(frame_spacer_01, width = 150, height = 110, background = colour)
    frame_8.grid(row=0, column=1, sticky = 'wn')
    frame_8.grid_propagate (False)
    
    laser= tk.Label(page3, text=' Laser:', height = 1, foreground= txtcolour, background=colour)
    laser.grid(row = 0, column = 0)
    laser_01 = tk.Entry(page3,width = 8)
    laser_01.insert(tk.END, '561')
    laser_01.grid(row = 0,column = 1)
    
    
#### Overview tab ######   

    space= tk.Label(page1, text='', height = 1 , foreground= txtcolour, background=colour)
    space.grid(row = 0, column = 0 , sticky = tk.W+tk.N)
    
    dwell_overview= tk.Label(page1, text=' Dwelltime [us]:', height = 1 , foreground= txtcolour, background=colour)
    dwell_overview.grid(row = 1, column = 0 , sticky = tk.W + tk.N)
    dwell_overview_value = tk.Entry(page1,width = 3, foreground= 'white', bg = 'grey')
    dwell_overview_value.insert(tk.END, '10')
    dwell_overview_value.grid(row = 1, column = 1, sticky = 'w')
    
    pxsize_overview= tk.Label(page1, text=' Pixelsize [nm]:', height = 1,foreground= txtcolour, background=colour)
    pxsize_overview.grid(row = 2, column = 0, sticky = 'wn')
    pxsize_overview_value = tk.Entry(page1, width = 3, foreground= 'white', bg = 'grey')
    pxsize_overview_value.insert(tk.END, '50')
    pxsize_overview_value.grid(row = 2, column = 1, sticky = 'w')
    
    ROIsize_overview= tk.Label(page1, text=' ROIsize [um]:', height = 1,foreground= txtcolour, background=colour)
    ROIsize_overview.grid(row = 3, column = 0, sticky = 'wn')
    ROIsize_overview_value = tk.Entry(page1,width = 3, foreground= 'white', bg = 'grey')
    ROIsize_overview_value.insert(tk.END, '10')
    ROIsize_overview_value.grid(row = 3, column = 1, sticky = 'w')
    
    frames_overview= tk.Label(page1, text=' # Frames:           ', height = 1, foreground= txtcolour, background=colour)
    frames_overview.grid(row = 4, column = 0, sticky = 'wn')
    frames_overview_value = tk.Entry(page1,width = 3, foreground= 'white', bg = 'grey')
    frames_overview_value.insert(tk.END, '3')
    frames_overview_value.grid(row = 4, column = 1, sticky = 'w')

    laser_overview= tk.Label(page1, text=' Laser| power[%]:', height = 1, foreground= txtcolour, background=colour)
    laser_overview.grid(row = 5, column = 0, sticky = 'wn')
    laser_overview_entry = tk.Entry(page1,width = 8, foreground= 'white', bg = 'grey')
    laser_overview_entry.insert(tk.END, '485')
    laser_overview_entry.grid(row = 5,column = 1, sticky = 'w') 
    
    laser_overview_value = tk.Entry(page1, width = 8, foreground= 'white', bg = 'grey')
    laser_overview_value.insert(tk.END, '5')
    laser_overview_value.grid(row = 5, column = 2, sticky = 'w')
    
    #this button should ideally be moved to a findpeak tab
    peakchannelLab= tk.Label(page1, text=' findpeak chan:', height = 1, foreground= txtcolour, background=colour)
    peakchannelLab.grid(row = 6, column = 0, sticky = 'w')
    peakchannels = tk.Entry(page1, width = 8, foreground= 'white', bg = 'grey')
    peakchannels.insert(tk.END, '0123')
    peakchannels.grid(row = 6, column = 1, sticky = 'w')

    
#########
### ROI_select 

    
    laservalue= tk.Label(page3, text=' value:', height = 1, foreground= txtcolour, background=colour)
    laservalue.grid(row = 0, column = 2, sticky = 'w')
    laser_value_01 = tk.Entry(page3,width = 8)
    laser_value_01.insert(tk.END, '50')
    laser_value_01.grid(row = 0, column = 3, sticky = 'w')
    
    laser_STED= tk.Label(page3, text=' STED-Laser:', height = 1, foreground= txtcolour, background=colour)
    laser_STED.grid(row = 1, column = 0, sticky = 'w')
    laser_STED_01 = tk.Entry(page3,width = 8)
    laser_STED_01.insert(tk.END, '775')
    laser_STED_01.grid(row = 1, column = 1, sticky = 'w')
    
    laser_STEDvalue= tk.Label(page3, text=' value:', height = 1, foreground= txtcolour, background=colour)
    laser_STEDvalue.grid(row = 1, column = 2)
    laser_STEDvalue_01 = tk.Entry(page3,width = 20)
    laser_STEDvalue_01.insert(tk.END, '50, 50, 20,32')
    laser_STEDvalue_01.grid(row = 1, column = 3)
    
    dwell= tk.Label(frame_5, text=' Dwelltime [us]:', height = 1 , foreground= txtcolour, background=colour)
    dwell.grid(row = 0, column = 0 , sticky = 'wn')
    dwell_01 = tk.Entry(frame_5,width = 3 )
    dwell_01.insert(tk.END, '5')
    dwell_01.grid(row = 0, column = 1, sticky = 'w')
    
    pxsize= tk.Label(frame_5, text=' Pixelsize [nm]:', height = 1,foreground= txtcolour, background=colour)
    pxsize.grid(row = 1, column = 0, sticky = 'wn')
    pxsize_01 = tk.Entry(frame_5, width = 3)
    pxsize_01.insert(tk.END, '10')
    pxsize_01.grid(row = 1, column = 1, sticky = 'w')
    
    ROIsize= tk.Label(frame_5, text=' ROIsize [um]:', height = 1,foreground= txtcolour, background=colour)
    ROIsize.grid(row = 2, column = 0, sticky = 'wn')
    ROIsize_01 = tk.Entry(frame_5,width = 3)
    ROIsize_01.insert(tk.END, '1')
    ROIsize_01.grid(row = 2, column = 1, sticky = 'w')
    
    frames= tk.Label(frame_5, text=' # Frames:           ', height = 1, foreground= txtcolour, background=colour)
    frames.grid(row = 3, column = 0, sticky = 'wn')
    frames_01 = tk.Entry(frame_5,width = 3)
    frames_01.insert(tk.END, '11')
    frames_01.grid(row = 3, column = 1, sticky = 'w')
    
    
    MultiRUN= tk.Label(frame_8, text=' Focus search:', height = 1, foreground= txtcolour, background=colour)
    MultiRUN.grid(row = 0, column = 2, sticky = 'w')
    MultiRUN_01 = tk.Entry(frame_8,width = 8)
    MultiRUN_01.insert(tk.END, '3')
    MultiRUN_01.grid(row = 0, column = 3, sticky = 'w')

    MultiRUN_s= tk.Label(frame_8, text='', height = 1, foreground= txtcolour, background=colour)
    MultiRUN_s.grid(row = 0, column = 0, sticky = 'w')
    
    Autofocus= tk.Label(frame_8, text=' Autofocus:', height = 1 ,foreground= txtcolour, background=colour)
    Autofocus.grid(row = 2, column = 2, sticky = tk.W)
    var13 = tk.IntVar()
    Autofocus_01 = tk.Checkbutton(frame_8,  variable = var13, background =colour)
    Autofocus_01.grid(row=2, column = 3, sticky = tk.W)
    
    QFS= tk.Label(page1, text=' Autofocus:', height = 1 ,foreground= txtcolour, background=colour)
    QFS.grid(row = 2, column = 2, sticky = tk.W)
    var14 = tk.IntVar()
    QFS_01 = tk.Checkbutton(page1,  variable = var14, background =colour)
    QFS_01.grid(row=2, column = 3, sticky = tk.W)
    
    Circle= tk.Label(frame_8, text=' Circle:', height = 1 ,foreground= txtcolour, background=colour)
    Circle.grid(row = 3, column = 2, sticky = tk.W)
    var15 = tk.IntVar(value=0)
    Circle_01 = tk.Checkbutton(frame_8,  variable = var15, background =colour)
    Circle_01.grid(row=3, column = 3, sticky = tk.W)
    

    
    
    L485= tk.Label(frame_6, text=' L485:', height = 1 ,foreground= txtcolour, background=colour)
    L485.grid(row = 0, column = 0, sticky = tk.W)
    var1 = tk.IntVar()
    L485_01 = tk.Checkbutton(frame_6,  variable = var1, background =colour)
    L485_01.grid(row=0, column = 1, sticky = tk.W)
    
    L485_02= tk.Label(frame_6, text=' L485:', height = 1 ,foreground= txtcolour, background=colour)
    L485_02.grid(row = 0, column = 0, sticky = tk.W)
    var7 = tk.IntVar()
    L485_02_ = tk.Checkbutton(frame_6,  variable = var7, background =colour)
    L485_02_.grid(row=0, column = 2, sticky=tk.W)
    
    L518= tk.Label(frame_6, text=' L518:', height = 1 ,foreground= txtcolour, background=colour)
    L518.grid(row = 1, column = 0, sticky = 'wn')
    var2 = tk.IntVar()
    L518_01 = tk.Checkbutton(frame_6,  variable = var2, background =colour)
    L518_01.grid(row=1, column = 1, sticky=tk.W)
    
    L518_02= tk.Label(frame_6, text=' L518:', height = 1,foreground= txtcolour, background=colour)
    L518_02.grid(row = 1, column = 0, sticky = 'wn')
    var8 = tk.IntVar()
    L518_02_ = tk.Checkbutton(frame_6,  variable = var8, background =colour)
    L518_02_.grid(row=1, column = 2, sticky=tk.W)
    
    L561= tk.Label(frame_6, text=' L561:', height = 1,foreground= txtcolour, background=colour)
    L561.grid(row = 2, column = 0, sticky = 'wn')
    var3 = tk.IntVar()
    L561_01 = tk.Checkbutton(frame_6,  variable = var3, background =colour)
    L561_01.grid(row=2, column = 1, sticky=tk.W)
    
    L561_02= tk.Label(frame_6, text=' L561:', height = 1,foreground= txtcolour, background=colour)
    L561_02.grid(row = 2, column = 0, sticky = 'wn')
    var9 = tk.IntVar()
    L561_02_ = tk.Checkbutton(frame_6,  variable = var9, background =colour)
    L561_02_.grid(row=2, column = 2, sticky=tk.W)
    
    L640= tk.Label(frame_6, text=' L640:', height = 1,foreground= txtcolour, background=colour)
    L640.grid(row = 3, column = 0, sticky = 'wn')
    var4 = tk.IntVar()
    L640_01 = tk.Checkbutton(frame_6, variable = var4, background =colour)
    L640_01.grid(row=3, column = 1, sticky=tk.W)
    
    L640_02= tk.Label(frame_6, text=' L640:', height = 1,foreground= txtcolour, background=colour)
    L640_02.grid(row = 3, column = 0, sticky = 'wn')
    var10 = tk.IntVar()
    L640_02_ = tk.Checkbutton(frame_6,  variable = var10, background =colour)
    L640_02_.grid(row=3, column = 2, sticky = 'wn')
    
    L595= tk.Label(frame_6, text=' L595:', height = 1,foreground= txtcolour, background=colour)
    L595.grid(row = 4, column = 0, sticky = 'wn')
    var5 = tk.IntVar()
    L595_01 = tk.Checkbutton(frame_6,  variable = var5, background =colour)
    L595_01.grid(row=4, column = 1, sticky = 'wn')
    
    L595_02= tk.Label(frame_6, text=' L595:', height = 1,foreground= txtcolour, background=colour)
    L595_02.grid(row = 4, column = 0, sticky = 'wn')
    var11 = tk.IntVar()
    L595_02_ = tk.Checkbutton(frame_6,  variable = var11, background =colour)
    L595_02_.grid(row=4, column = 2, sticky = 'wn')
    
    L775= tk.Label(frame_6, text=' L775:', height = 1,foreground= txtcolour, background=colour)
    L775.grid(row = 5, column = 0, sticky = 'wn')
    var6 = tk.IntVar()
    L775_01 = tk.Checkbutton(frame_6, variable = var6, background =colour)
    L775_01.grid(row=5, column = 1, sticky = 'wn')
    
    L775_02= tk.Label(frame_6, text=' L775:', height = 1,foreground= txtcolour, background=colour)
    L775_02.grid(row = 5, column = 0, sticky = 'wn')
    var12 = tk.IntVar()
    L775_02_ = tk.Checkbutton(frame_6,  variable = var12, background =colour)
    L775_02_.grid(row=5, column = 2, sticky = 'wn')
    
    L485_value= tk.Label(frame_6, height = 1,foreground= txtcolour, background=colour)
    L485_value.grid(row = 2, column = 0, sticky = 'wn')
    L485_value_01 = tk.Entry(frame_6,width = 3)
    L485_value_01.insert(tk.END, '0')
    L485_value_01.grid(row = 0, column = 5, sticky = 'en')
    
    L518_value= tk.Label(frame_6, height = 1,foreground= txtcolour, background=colour)
    L518_value.grid(row = 1, column = 0, sticky = 'wn')
    L518_value_01 = tk.Entry(frame_6,width = 3)
    L518_value_01.insert(tk.END, '0')
    L518_value_01.grid(row = 1, column = 5, sticky = 'en')
    
    L561_value= tk.Label(frame_6, height = 1,foreground= txtcolour, background=colour)
    L561_value.grid(row = 2, column = 0, sticky = 'wn')
    L561_value_01 = tk.Entry(frame_6,width = 3)
    L561_value_01.insert(tk.END, '50')
    L561_value_01.grid(row = 2, column = 5, sticky = 'en')
    
    L640_value= tk.Label(frame_6, height = 1,foreground= 'black', background=colour)
    L640_value.grid(row = 3, column = 0, sticky = 'wn')
    L640_value_01 = tk.Entry(frame_6,width = 3)
    L640_value_01.insert(tk.END, '3')
    L640_value_01.grid(row = 3, column = 5, sticky = 'en')
    
    L595_value= tk.Label(frame_6, height = 1,foreground= 'black', background= colour)
    L595_value.grid(row = 4, column = 0, sticky = 'wn')
    L595_value_01 = tk.Entry(frame_6,width = 3)
    L595_value_01.insert(tk.END, '0')
    L595_value_01.grid(row = 4, column = 5, sticky = 'en')
    
    L775_value= tk.Label(frame_6, height = 1,foreground= 'black', background=colour)
    L775_value.grid(row = 5, column = 0, sticky = 'wn')
    L775_value_01 = tk.Entry(frame_6,width = 3)
    L775_value_01.insert(tk.END, '50')
    L775_value_01.grid(row = 5, column = 5, sticky = 'en')
    
    THRES_value= tk.Label(frame_6, height = 1,foreground= 'black', background=colour)
    THRES_value.grid(row = 6, column = 0, sticky = 'wn')
    THRES_value_01 = tk.Entry(frame_6,width = 3)
    THRES_value_01.insert(tk.END, '0')
    THRES_value_01.grid(row = 6, column = 5, sticky = 'en')
    
    a=0
    #put variables from GUI in class variables such that they can be accessed
    #from anywhere
    #sometimes the naming is changed for better readability
    self.frame_topleft = frame_top
    self.frame_botleft = frame_top2
    self.frame_topright = frame_top3
    self.frame_botright = frame_top4
    self.frame_buttons = labtext_1
    self.T = T
    self.frame_spacer_01 = frame_spacer_01 # shouldn't need
    self.frame_5 = frame_5 # shouldn't need
    self.frame_6 = frame_6#Shouldn't need
    self.frame_7 = frame_7 # shouldn't need
    self.frame_8 = frame_8 # shouldn't need
    self.color = colour # move to american spelling # shouldn't need

    self.pxsize = pxsize_01
    self.ROIsize = ROIsize_01
    self.dwelltime = dwell_01
    self.NoFrames = frames_01
    self.L485_1 = var1
    self.L518_1 = var2
    self.L561_1 = var3
    self.L640_1 = var4
    self.L595_1 = var5
    self.L775_1 = var6
    self.L485_2 = var7
    self.L518_2 = var8
    self.L561_2 = var9
    self.L640_2 = var10
    self.L595_2 = var11
    self.L775_2 = var12
    self.AutofocusOnROISelect = var13
    self.AutofocusOnOverview = var14
    self.circle = var15
    self.peakchannels = peakchannels
    self.L485_value = L485_value_01 #don't need this 01 appendix
    self.L518_value = L518_value_01
    self.L561_value = L561_value_01
    self.L640_value = L640_value_01
    self.L595_value = L595_value_01
    self.L775_value = L775_value_01
    self.multirun = MultiRUN_01
    self.laser_overview_value = laser_overview_value
    self.laser_overview_entry = laser_overview_entry
    self.frames_overview_value = frames_overview_value
    self.ROIsize_overview_value = ROIsize_overview_value
    self.dwell_overview_value = dwell_overview_value
    self.pxsize_overview_value = pxsize_overview_value
    return    





























